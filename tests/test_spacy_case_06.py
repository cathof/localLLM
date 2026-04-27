#!/usr/bin/env python3
from pathlib import Path
from typing import Any, Dict, List
import re


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts: List[str] = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def split_document_into_segments(
        text: str,
        target_chars: int = 1200,
        max_chars: int = 2200,
) -> List[str]:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", text) if p.strip()]

    segments: List[str] = []
    current: List[str] = []

    def flush() -> None:
        nonlocal current
        if current:
            segments.append("\n\n".join(current).strip())
            current = []

    for para in paragraphs:
        candidate = "\n\n".join(current + [para]).strip() if current else para

        if len(candidate) > target_chars and current:
            flush()

        if len(para) > max_chars:
            chunks = [para[i:i + max_chars] for i in range(0, len(para), max_chars)]
            for c in chunks:
                segments.append(c.strip())
            continue

        current.append(para)

    flush()
    return segments


def run_spacy_adjective_check(segment_text: str, segment_index: int) -> List[Dict[str, Any]]:
    import spacy

    if not hasattr(run_spacy_adjective_check, "_nlp"):
        run_spacy_adjective_check._nlp = spacy.load("de_dep_news_trf")  # type: ignore[attr-defined]
        print("[INFO] spaCy de_dep_news_trf loaded")

    nlp = run_spacy_adjective_check._nlp  # type: ignore[attr-defined]
    doc = nlp(segment_text)

    findings: List[Dict[str, Any]] = []

    for token in doc:
        if token.text and token.text[0].isupper() and token.pos_ == "ADJ":
            prev_pos = doc[token.i - 1].pos_ if token.i > 0 else "START"
            print(
                f"[DEBUG-SPACY] S{segment_index} "
                f"ADJ={token.text!r} dep={token.dep_} "
                f"head={token.head.text!r}({token.head.pos_}) prev={prev_pos}"
            )

        if (
                token.pos_ == "ADJ"
                and token.dep_ == "nk"
                and token.head.pos_ == "NOUN"
                and token.text
                and token.text[0].isupper()
                and token.i > 0
                and doc[token.i - 1].pos_ == "DET"
        ):
            correct = token.text[0].lower() + token.text[1:]
            context = doc[max(0, token.i - 3): token.i + 4].text

            findings.append({
                "segment_index": segment_index,
                "stelle_im_segment": token.text,
                "vorschlag": correct,
                "context": context,
                "dep": token.dep_,
                "head": token.head.text,
                "head_pos": token.head.pos_,
                "prev": doc[token.i - 1].text,
                "prev_pos": doc[token.i - 1].pos_,
            })

    return findings


def main() -> None:
    path = Path("./case_documents/case_06.docx")

    if not path.exists():
        raise SystemExit(f"Document not found: {path.resolve()}")

    text = read_docx(path)
    print(f"[INFO] Document chars: {len(text)}")

    segments = split_document_into_segments(text)
    print(f"[INFO] Segments: {len(segments)}")

    all_findings: List[Dict[str, Any]] = []

    for i, segment in enumerate(segments, start=1):
        findings = run_spacy_adjective_check(segment, i)
        all_findings.extend(findings)

    print("\n" + "=" * 80)
    print("SPACY FINDINGS")
    print("=" * 80)

    if not all_findings:
        print("Keine spaCy-Findings nach aktueller Regel gefunden.")
        return

    for n, f in enumerate(all_findings, start=1):
        print(f"\n{n}. Segment {f['segment_index']}")
        print(f"   Stelle:    {f['stelle_im_segment']}")
        print(f"   Vorschlag: {f['vorschlag']}")
        print(f"   Kontext:   {f['context']}")
        print(
            f"   Analyse:   dep={f['dep']} "
            f"head={f['head']}({f['head_pos']}) "
            f"prev={f['prev']}({f['prev_pos']})"
        )


if __name__ == "__main__":
    main()