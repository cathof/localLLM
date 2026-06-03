#!/usr/bin/env python3
"""
synthetic_error_injector.py
============================
Generiert synthetische Fehlervorschläge für ein .docx-Gutachten.
Fehlerklassen abgeleitet aus Ground Truth Case 06 + Taxonomie v1.

UNIVERSELLES MUSTER (pro Klasse):
  Phase 1 – Kandidatenextraktion  (Regex oder leichter LLM-Call)
  Phase 2 – Budget-Selektor       (Python, klassenübergreifend)
  Phase 3 – Injektion             (deterministisch oder LLM)

FEHLERKLASSEN (aus GT, nach Häufigkeit):
  STRUKT_BEFUND_BESCHREIBUNG  36 %  LLM       MESSWERT/INSTITUTION/PERSON/DATUM
  FORMALES_REDAKTION          29 %  LLM       Zufälliger Subtyp (Tippfehler, Gross-/Kleinschreibung, Doppelkonsonant etc.)
  RECHT_ABSICHERUNG           14 %  LLM       Hedging entfernen → zu absolut
  RECHEN_ARITHMETIK           14 %  determin. Rechenoperand perturbieren
  HYPO_INKONSISTENZ            4 %  LLM       Hypothesenbewertung kippen
  RECHT_TRENNUNG               4 %  LLM       Konjunktiv → Indikativ

WORKFLOW (4 Schritte):
  1. python synthetic_error_injector.py --mode extract_facts \\
         --document ./case_documents/case_02.docx
     → injection/reference_facts_case_02_original.json

  2. python synthetic_error_injector.py --mode generate \\
         --document ./case_documents/case_02.docx
     → injection/proposals_case_02_latest.json
     (proposals_*.json öffnen, status="accepted"/"rejected" setzen)

  3. python synthetic_error_injector.py --mode inject \\
         --document ./case_documents/case_02.docx \\
         --proposals injection/proposals_case_02_latest.json
     → case_documents/case_02_modified.docx
     (Kopie des Originaldokuments mit injizierten Fehlern)

  4. python synthetic_error_injector.py --mode validate \\
         --proposals injection/proposals_case_02_latest.json \\
         --modified_doc ./case_documents/case_02_modified.docx
     → injection/ground_truth_case_02_synthetic_<ts>.jsonl

ENV-VARIABLEN (.env):
  OLLAMA_BASE_URL, LLM_BACKEND, LLM_TIMEOUT_S, CASE_ID, TAXONOMY_JSON,
  REFERENCE_FACTS_SCHEMA, REFERENCE_FACTS_CONTEXT_CHARS,
  INJECTION_MODEL, INJECTION_TEMPERATURE (default 0.3),
  INJECTION_TARGET_ERRORS (default 10), INJECTION_SEED (default 42),
  INJECTION_OUTPUT_DIR (default injection),
  INJECTION_ENTITY_WHITELIST (JSON-Array Schweizer Institutionen),
  SEG_TARGET_CHARS, SEG_MIN_CHARS, SEG_MAX_CHARS
"""

from __future__ import annotations

import argparse
import json
import random
import re
import shutil
import sys
import time
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


# ── Bestehende Infrastruktur ──────────────────────────────────────────────────
_THIS_DIR = Path(__file__).resolve().parent
if str(_THIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THIS_DIR))

try:
    from rag_answer_reference_facts_fp_hoch import (
        load_dotenv, env_str, env_int, env_bool,
        require_env, env_json_object_optional,
        load_taxonomy_json, ErrorCatalog,
        LLMClient, OllamaClient,
        split_document_into_segments,
        run_reference_facts_agent,
        load_reference_facts_schema,
        normalize_reference_facts,
        enrich_reference_facts_from_document,
        extract_first_json_object,
    )
except ImportError as e:
    sys.exit(f"[FATAL] rag_answer_reference_facts.py Importfehler:\n  {e}")

try:
    from importDocuments_structural import normalize_text, read_docx
except ImportError as e:
    sys.exit(f"[FATAL] importDocuments_structural.py nicht gefunden:\n  {e}")

load_dotenv(".env")


# ═══════════════════════════════════════════════════════════════════════════════
# KONFIGURATION  –  abgeleitet aus Ground Truth Case 06 (28 Findings)
# ═══════════════════════════════════════════════════════════════════════════════

# Zielverteilung aus GT (Anteil an allen Findings)
CLASS_BUDGET: Dict[str, float] = {
    "STRUKT_BEFUND_BESCHREIBUNG": 0.36,   # 10/28
    "FORMALES_REDAKTION":         0.29,   # 8/28
    "RECHT_ABSICHERUNG":          0.14,   # 4/28
    "RECHEN_ARITHMETIK":          0.14,   # 4/28
    "HYPO_INKONSISTENZ":          0.04,   # 1/28
    "RECHT_TRENNUNG":             0.04,   # 1/28
}

# Taxonomie-Metadaten je Klasse (aus GT-Findings)
CLASS_META: Dict[str, Dict[str, str]] = {
    "STRUKT_BEFUND_BESCHREIBUNG": {
        "change_type_id": "CHANGE_FACHLICH",
        "severity_id":    "HIGH",
    },
    "FORMALES_REDAKTION": {
        "change_type_id": "CHANGE_REDAKTION",
        "severity_id":    "LOW",
    },
    "RECHT_ABSICHERUNG": {
        "change_type_id": "CHANGE_FACHLICH",
        "severity_id":    "MEDIUM",
    },
    "RECHEN_ARITHMETIK": {
        "change_type_id": "CHANGE_RECHNERISCH",
        "severity_id":    "HIGH",
    },
    "HYPO_INKONSISTENZ": {
        "change_type_id": "CHANGE_HYPOTHESE",
        "severity_id":    "MEDIUM",
    },
    "RECHT_TRENNUNG": {
        "change_type_id": "CHANGE_ARGUMENTATION",
        "severity_id":    "MEDIUM",
    },
}

# Reihenfolge = Priorität bei Budget-Gleichstand
CLASS_PRIORITY: List[str] = [
    "STRUKT_BEFUND_BESCHREIBUNG",
    "FORMALES_REDAKTION",
    "RECHEN_ARITHMETIK",
    "RECHT_ABSICHERUNG",
    "RECHT_TRENNUNG",
    "HYPO_INKONSISTENZ",
]

# Schweizer Entitäten-Whitelist (Fallback)
_DEFAULT_ENTITY_WHITELIST: List[str] = [
    "Staatsanwaltschaft Zürich-Sihl",
    "Staatsanwaltschaft Zürich-Limmat",
    "Staatsanwaltschaft See/Oberland",
    "Staatsanwaltschaft Winterthur/Unterland",
    "Staatsanwaltschaft Zürich-Oerlikon",
    "Staatsanwaltschaft Baden",
    "Staatsanwaltschaft Brugg-Zurzach",
    "Staatsanwaltschaft Bern-Mittelland",
    "Staatsanwaltschaft Luzern",
    "Kantonspolizei Zürich",
    "Kantonspolizei Bern",
    "Kantonspolizei Aargau",
    "Kantonspolizei Luzern",
    "Kantonspolizei St. Gallen",
    "Kantonsspital Winterthur",
    "Kantonsspital Aarau",
    "Universitätsspital Zürich",
    "Inselspital Bern",
    "Institut für Rechtsmedizin Zürich",
    "Institut für Rechtsmedizin Bern",
    "Forensisch-Medizinischer Dienst Zürich",
    "Bezirksgericht Winterthur",
    "Bezirksgericht Zürich",
    "Obergericht des Kantons Zürich",
]


# ═══════════════════════════════════════════════════════════════════════════════
# KANDIDATEN-DATENSTRUKTUR
# ═══════════════════════════════════════════════════════════════════════════════
#
# Jeder Eintrag aus Phase 1 (Kandidatenextraktion) hat dieses Format:
#
# {
#   "span_text":             str,   # Originaltext, der ersetzt wird
#   "replacement":           str,   # Vorgefertigter Ersatz (determin.) oder "" (LLM)
#   "subtype":               str,   # Subtyp für Logging (FREQUENZ, TYPO, …)
#   "is_reformulation":      bool,  # True = ganzer Satz wird umgeschrieben
#   "requires_human_review": bool,  # True = zwingend manuelle Prüfung vor GT
#   "rationale":             str,   # Vorläufige Begründung (determin. Injektoren)
# }


# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 1 – KANDIDATENEXTRAKTION  (pro Fehlerklasse)
# ═══════════════════════════════════════════════════════════════════════════════

# ── STRUKT_BEFUND_BESCHREIBUNG ────────────────────────────────────────────────

_MEASUREMENT_RE = re.compile(
    r"\d+[.,]?\d*\s*"
    r"(?:Hz|kHz|MHz|GHz|"
    r"Meter|Metern|m(?=\b)|km\b|km/h|"
    r"kg\b|Kilogramm|g(?=\b)|mg\b|"
    r"ml\b|Liter|l(?=\b)|"
    r"Promille|%(?=\b)|"
    r"cm\b|mm\b|m/s|km²|m²|"
    r"kPa|bar|°C|Pa|dB|dBFS|dBSPL|"
    r"Joule|kJ|MJ|Watt|kW|Bit\b|"
    r"Sekunden|Minuten|Stunden|ms\b)",
    re.IGNORECASE,
)
_DATE_RE = re.compile(
    r"\b\d{1,2}\.\s*"
    r"(?:Januar|Februar|März|April|Mai|Juni|Juli|August|"
    r"September|Oktober|November|Dezember)\s*\d{4}",
    re.IGNORECASE,
)
# Matcht nur den Keyword-Anker – wird intern für Kandidatenfilterung genutzt
_INSTITUTION_RE = re.compile(
    r"\b(?:Staatsanwaltschaft|Kantonspolizei|Kantonsgericht|"
    r"Bezirksgericht|Obergericht|Bundesgericht|Spital|Klinik|"
    r"Institut\s+(?:für|of)|Forensisch|IRM\b|Polizei\b|"
    r"Universität|Fachhochschule|ZHAW|ETH\b)",
    re.IGNORECASE,
)
# Extrahiert den vollständigen Institutionsnamen (Keyword + Kontext)
_INSTITUTION_FULL_RE = re.compile(
    r"(?:"
    r"Staatsanwaltschaft\s+[\w/\-]+(?:\s+[\w/\-]+)?"
    r"|Kantonspolizei\s+\w+(?:\s+\w+)?"
    r"|(?:Kantons|Bezirks|Ober)gericht\s+\w+(?:\s+\w+)?"
    r"|(?:Kantons|Universitäts|Insel)spital\s+\w+"
    r"|Institut\s+für\s+\w+(?:\s+\w+){0,3}"
    r"|Forensisch(?:es|er|en|em)?\s+(?:Medizinischer\s+)?(?:Institut|Dienst|Labor)\s+\w+(?:\s+\w+)?"
    r"|IRM\s+\w+(?:\s+\w+)?"
    r"|Bundesgericht\b"
    r"|ZHAW\b"
    r"|ETH\s+\w+"
    r")",
    re.IGNORECASE,
)
_PERSON_TITLE_RE = re.compile(r"\b(?:Dr\.|Prof\.|dipl\.|lic\.|MLaw)", re.IGNORECASE)


def extract_candidates_strukt(
        segment_text: str,
        reference_facts: Optional[Dict[str, Any]] = None,
) -> Dict[str, List[Dict]]:
    """
    Gibt dict {subtype: [candidates]} zurück, je mit konkretem span_text.
    Phase 1 extrahiert echte Spans – das LLM muss in Phase 3 nur noch ersetzen,
    nicht mehr suchen. Das verhindert semantische Fehlklassifikationen.
    subtype ∈ {MESSWERT, INSTITUTION, PERSON, DATUM}
    """
    result: Dict[str, List[Dict]] = {
        "MESSWERT": [], "INSTITUTION": [], "PERSON": [], "DATUM": []
    }

    # MESSWERT: alle (Zahl + Einheit)-Matches als konkrete Spans
    for m in _MEASUREMENT_RE.finditer(segment_text):
        span = m.group(0).strip()
        if span:
            result["MESSWERT"].append({
                "span_text": span, "replacement": "", "subtype": "MESSWERT",
                "is_reformulation": False, "requires_human_review": False, "rationale": "",
            })

    # DATUM: alle vollständigen Datumsangaben
    for m in _DATE_RE.finditer(segment_text):
        span = m.group(0).strip()
        if span:
            result["DATUM"].append({
                "span_text": span, "replacement": "", "subtype": "DATUM",
                "is_reformulation": False, "requires_human_review": False, "rationale": "",
            })

    # INSTITUTION: vollständige Institutionsnamen via erweiterter Regex.
    # Kein Fallback auf Keyword-Anker: ein einzelnes Keyword wie "Forensisch"
    # ist kein sinnvoller Span (kann Adjektiv oder Teil eines Titels sein).
    # Guillemet-Filter: Spans innerhalb «» sind Dokumenttitel, keine Institutionen.
    seen_inst: set = set()
    for m in _INSTITUTION_FULL_RE.finditer(segment_text):
        span = m.group(0).strip()
        if not span or span in seen_inst:
            continue
        # Prüfen ob Span innerhalb «...» liegt
        preceding = segment_text[:m.start()]
        if preceding.count("«") > preceding.count("»"):
            continue
        seen_inst.add(span)
        result["INSTITUTION"].append({
            "span_text": span, "replacement": "", "subtype": "INSTITUTION",
            "is_reformulation": False, "requires_human_review": False, "rationale": "",
        })

    # PERSON: nur bekannte Referenzfakten-Namen die im Segment vorkommen
    if reference_facts:
        known = _known_person_names(reference_facts)
        for name in known:
            if name in segment_text:
                result["PERSON"].append({
                    "span_text": name, "replacement": "", "subtype": "PERSON",
                    "is_reformulation": False, "requires_human_review": False, "rationale": "",
                })

    return result


# ── FORMALES_REDAKTION  (vollständig deterministisch) ─────────────────────────

# Deutsche Verwechslungen für realistische Tippfehler
# Subtypen für FORMALES_REDAKTION – zufällig gewählt, kein Budget.
# Diversität entsteht durch Zufall, nicht durch Überanpassung an die GT.
_FORMALES_SUBTYPES: List[str] = [
    "TIPPFEHLER",           # Buchstaben getauscht/gelöscht     (z.B. "Kollison")
    "DOPPELKONSONANT",      # Doppelkonsonant reduziert          (z.B. "Strase")
    "ADJEKTIV_GROSS",       # Adjektiv fälschlich grossgeschrieben
    "HOMOPHON",             # Homophon-Verwechslung              (z.B. "das" → "dass")
    "FREMDWORT_FALSCH",     # Fremdwort falsch geschrieben       (z.B. "Analyse" → "Analise")
    # KOMPOSITA_TRENNUNG entfernt: LanguageTool erkennt valide Einzelwörter nicht
    # als Fehler → strukturell nicht erkennbar → immer FN
]

# Beschreibungen je Subtyp – eingebettet im LLM-Prompt
_FORMALES_SUBTYPE_DESC: Dict[str, str] = {
    "TIPPFEHLER": (
        "Erzeuge einen klassischen Tippfehler: Tausche zwei benachbarte Buchstaben "
        "oder lösche einen einzelnen Buchstaben im Wortinnern. "
        "Beispiel: 'Kollision' → 'Kolilsion' oder 'Kolision'."
    ),
    "DOPPELKONSONANT": (
        "Entferne einen Doppelkonsonanten: Reduziere 'ss', 'tt', 'll', 'nn', 'mm', 'pp' etc. "
        "auf einen einfachen Konsonanten. "
        "Beispiel: 'Strasse' → 'Strase', 'Kollision' → 'Kolision'."
    ),
    "ADJEKTIV_GROSS": (
        "Schreibe ein attributives Adjektiv fälschlicherweise gross. "
        "Das Adjektiv muss direkt vor einem Substantiv stehen. "
        "Beispiel: 'die beschuldigte Person' → 'die Beschuldigte Person'."
    ),
    "HOMOPHON": (
        "Ersetze ein Wort durch ein gleichklingendes aber falsch geschriebenes Homophon.\n"
        "Verwende NUR Paare aus dieser Liste — keine anderen:\n"
        "  dass → das  |  das → dass  |  seit → seid  |  seid → seit\n"
        "  so dass → sodass  |  sodass → so dass\n"
        "  wieder → wider  |  wider → wieder\n"
        "  als → als (NICHT verwenden)\n"
        "Wähle das Paar das im Segment vorkommt und ersetze NUR dieses eine Wort.\n"
        "NICHT auf Eigennamen, Zahlen oder Abkürzungen anwenden."
    ),
    "FREMDWORT_FALSCH": (
        "Schreibe ein Fremd- oder Fachwort phonetisch falsch. "
        "Beispiel: 'Analyse' → 'Analise', 'Forensik' → 'Forensick', "
        "'Hypothese' → 'Hypotese'."
    ),
}


# Precheck-Regexes: prüfen ob ein Subtyp im Segment überhaupt anwendbar ist.
# Verhindert dass das Modell einen Subtyp halluziniert, der strukturell nicht passt.
_DOUBLE_CONSONANT_RE = re.compile(
    r'\b\w*(?:ss|tt|ll|nn|mm|pp|rr|gg|bb|dd|ff|kk|zz)\w*\b',
    re.IGNORECASE,
)
_LONG_COMPOUND_RE = re.compile(r'\b[A-ZÜÄÖ][a-züäöß]{8,}\b')
_FREMDWORT_RE     = re.compile(
    r'\b(?:Analys|Forensi|Hypothes|Methodi|Indizi|Expertis|Protokoll|'
    r'Kompress|Frequenz|Akustik|Digital|Spektr|Signal|Kalibrierung|'
    r'Parameter|Algorithmus|Korrelation)\w*\b',
    re.IGNORECASE,
)
_ADJEKTIV_RE      = re.compile(r'\b[a-züäöß]{5,}\b')   # Kleinwörter ≥5 als Proxy
def _find_double_consonant_words(text: str) -> List[str]:
    """Gibt alle Wörter zurück, die tatsächlich einen Doppelkonsonanten enthalten."""
    return [
        m.group(0) for m in _DOUBLE_CONSONANT_RE.finditer(text)
        if len(m.group(0)) >= 4
    ]


# Wörter mit echtem Umlaut (ä, ö, ü) – Grossbuchstaben-Varianten eingeschlossen,
# aber Eigennamen ausschliessen (beginnen mit Grossbuchstabe nach Satzanfang)
# Kleinbuchstaben-Wörter ≥5 Zeichen, keine Eigennamen (kein Grossbuchstabe innen)
_LOWER_WORD_RE = re.compile(r'\b[a-züäöß][a-züäöß]{4,}\b')

def _find_tippfehler_words(text: str) -> List[str]:
    """Gibt geeignete Wörter für Tippfehler zurück:
    Nur Kleinbuchstaben-Wörter (kein Eigenname), mind. 5 Buchstaben,
    keine Abkürzungen oder Zahlen.
    """
    return [
        m.group(0) for m in _LOWER_WORD_RE.finditer(text)
        if len(m.group(0)) >= 5
    ]


# Attributive Adjektive: Kleinwörter ≥5 Zeichen die vor einem Nomen stehen.
# Als Proxy: Wörter die auf -en, -em, -er, -es, -e enden (Adjektiv-Flexion)
_ADJEKTIV_WORD_RE = re.compile(
    r'\b[a-züäöß]{4,}(?:en|em|er|es|e)\b'
)

def _find_adjektiv_words(text: str) -> List[str]:
    """Gibt mögliche attributive Adjektive zurück (heuristisch über Flexionsendung)."""
    return [
        m.group(0) for m in _ADJEKTIV_WORD_RE.finditer(text)
        if len(m.group(0)) >= 5
    ]


_SUBTYPE_PRECHECK: Dict[str, Any] = {
    # Alle Prüfungen auf Wort-Ebene: nur Subtypen, bei denen mind. ein
    # geeignetes Wort im Segment nachweisbar ist, gelten als anwendbar.
    "DOPPELKONSONANT":    lambda t: bool(_find_double_consonant_words(t)),
    "ADJEKTIV_GROSS":     lambda t: bool(_find_adjektiv_words(t)),
    "HOMOPHON": lambda t: bool(re.search(
        r"\b(?:dass|das|seit|seid|sodass|so\s+dass|wieder|wider)\b", t, re.IGNORECASE
    )),
    "FREMDWORT_FALSCH":   lambda t: bool(_FREMDWORT_RE.search(t)),
    "TIPPFEHLER":         lambda t: bool(_find_tippfehler_words(t)),
}


# Mapping Subtyp → Wort-Extraktorfunktion
# Alle Subtypen mit dediziertem Extraktor liefern candidate_words an den Prompt,
# damit das Modell nicht frei im Segment suchen und dabei halluzinieren kann.
_SUBTYPE_WORD_EXTRACTOR: Dict[str, Any] = {
    "DOPPELKONSONANT":    _find_double_consonant_words,
    "TIPPFEHLER":         _find_tippfehler_words,
    "ADJEKTIV_GROSS":     _find_adjektiv_words,
    # HOMOPHON: deterministischer Extraktor — gibt das gefundene Homophon-Wort zurück
    "HOMOPHON": lambda t: [
        m.group(0) for m in re.finditer(
            r"\b(?:dass|das|seit|seid|sodass|so\s+dass|wieder|wider)\b",
            t, re.IGNORECASE
        )
    ][:3],
    # FREMDWORT: kein Wort-Extraktor, LLM sucht selbst
    "FREMDWORT_FALSCH":   None,
}


def extract_candidates_formales(
        segment_text: str,
        rng: random.Random,
) -> List[Dict]:
    """
    Gibt pro Segment eine Liste der anwendbaren Formales-Subtypen zurück.
    Für alle Subtypen mit Wort-Extraktor werden konkrete Kandidatenwörter
    mitgeliefert – das LLM darf nur noch aus diesen wählen.
    """
    candidates = []
    for st in _FORMALES_SUBTYPES:
        if not _SUBTYPE_PRECHECK[st](segment_text):
            continue
        entry: Dict[str, Any] = {"subtype": st, "candidate_words": []}
        extractor = _SUBTYPE_WORD_EXTRACTOR.get(st)
        if extractor is not None:
            words = extractor(segment_text)
            entry["candidate_words"] = list(dict.fromkeys(words))[:5]  # max 5, dedupliziert
        candidates.append(entry)
    return candidates


# ── RECHEN_ARITHMETIK  (vollständig deterministisch) ──────────────────────────

# Muster: "N Einheit (A – B Einheit)" oder "N Einheit (A + B Einheit)"
_CALC_DIFF_RE = re.compile(
    r"(\d+(?:[.,]\d+)?)\s+(\w+)\s*\("
    r"(\d+(?:[.,]\d+)?)\s+\w+\s*[–\-]\s*(\d+(?:[.,]\d+)?)\s+\w+\s*\)",
    re.IGNORECASE,
)
# Einfaches Muster: "A – B = C" oder "A - B = C"
_CALC_SIMPLE_RE = re.compile(
    r"(\d+(?:[.,]\d+)?)\s*[–\-]\s*(\d+(?:[.,]\d+)?)\s*[=]\s*(\d+(?:[.,]\d+)?)",
)
# Multiplikation: "A × B = C"
_CALC_MUL_RE = re.compile(
    r"(\d+(?:[.,]\d+)?)\s*[×\*]\s*(\d+(?:[.,]\d+)?)\s*[=]\s*(\d+(?:[.,]\d+)?)",
)


# Geschwindigkeit aus Weg/Zeit: d m / t s = v m/s → km/h
# Muster: "d m : t s = v m/s" oder "d m ÷ t s × 3.6 = V km/h"
_CALC_SPEED_DIV_RE = re.compile(
    r"(\d+(?:[.,]\d+)?)\s*m\s*[÷:/]\s*(\d+(?:[.,]\d+)?)\s*s"
    r"(?:\s*[×\*]\s*3[.,]6)?"
    r"\s*[=≈]\s*(\d+(?:[.,]\d+)?)\s*(km/h|m/s)",
    re.IGNORECASE,
)

# Durchschnittsgeschwindigkeit: "Strecke m in Zeit s → V km/h"
# Muster: "V km/h" nach explizitem Weg und Zeit im Satz
_CALC_SPEED_TEXT_RE = re.compile(
    r"(\d+(?:[.,]\d+)?)\s*m[^.!?]{0,80}"
    r"(\d+(?:[.,]\d+)?)\s*s[^.!?]{0,80}"
    r"(\d+(?:[.,]\d+)?)\s*km/h",
    re.IGNORECASE,
)

# Division: "A / B = C" (Weg/Zeit, allgemein)
_CALC_DIV_RE = re.compile(
    r"(\d+(?:[.,]\d+)?)\s*[:/]\s*(\d+(?:[.,]\d+)?)"
    r"\s*[=≈]\s*(\d+(?:[.,]\d+)?)",
)


def _trunc1(v: float) -> float:
    """Abschneiden auf 1 Nachkommastelle (Gutachten-Konvention)."""
    import math
    return math.floor(v * 10) / 10


def _parse_num(s: str) -> float:
    """Parst deutsche Zahl — unterscheidet Tausender-Punkt von Dezimal-Punkt.
    Tausender-Punkt: "25.320" (≥4 Stellen vor dem Punkt → 25320)
    Dezimal-Punkt:   "3.848" (nur 1 Stelle vor Punkt → 3.848)
    Komma als Dezimal: "3,848" → 3.848
    """
    s = s.strip()
    if "." in s and "," not in s:
        parts = s.split(".")
        # Tausender-Trenner nur wenn mind. 4 Ziffern vor dem Punkt
        if len(parts[0]) >= 4 and len(parts[-1]) == 3:
            return float(s.replace(".", ""))
    return float(s.replace(",", "."))


def _format_num(n: float, reference: str) -> str:
    """Formatiert Zahl analog zur Referenz."""
    ref = reference.strip()
    if "." in ref and len(ref.split(".")[-1]) == 3 and "," not in ref:
        return f"{int(n):,}".replace(",", ".")
    if "," in ref:
        decimals = len(ref.split(",")[-1])
        return f"{n:.{decimals}f}".replace(".", ",")
    if "." in ref:
        decimals = len(ref.split(".")[-1])
        return f"{n:.{decimals}f}"
    return str(int(round(n)))


def extract_candidates_rechen(segment_text: str) -> List[Dict]:
    candidates = []

    # Muster 1: "N Einheit (A – B Einheit)"
    for m in _CALC_DIFF_RE.finditer(segment_text):
        result_str, unit, a_str, b_str = m.group(1), m.group(2), m.group(3), m.group(4)
        try:
            a, b = _parse_num(a_str), _parse_num(b_str)
            result = _parse_num(result_str)
            correct = a - b
            if abs(result - correct) < max(0.1, abs(correct) * 0.02):
                # Dokument ist korrekt → Fehler injizieren
                delta = 1 if correct == int(correct) else correct * 0.11
                wrong = correct + delta
                wrong_str = _format_num(wrong, result_str)
                candidates.append({
                    "span_text":             result_str,
                    "replacement":           wrong_str,
                    "subtype":               "DIFFERENZ",
                    "is_reformulation":      False,
                    "requires_human_review": False,
                    "rationale": (
                        f"Korrekte Differenz {a_str}–{b_str}={_format_num(correct, result_str)} {unit}; "
                        f"injiziert: {wrong_str} {unit}"
                    ),
                })
        except (ValueError, ZeroDivisionError):
            continue

    # Muster 2: "A – B = C"
    for m in _CALC_SIMPLE_RE.finditer(segment_text):
        a_str, b_str, result_str = m.group(1), m.group(2), m.group(3)
        try:
            a, b, result = _parse_num(a_str), _parse_num(b_str), _parse_num(result_str)
            correct = a - b
            if abs(result - correct) < max(0.1, abs(correct) * 0.02):
                delta = 1 if correct == int(correct) else correct * 0.11
                wrong = correct + delta
                wrong_str = _format_num(wrong, result_str)
                candidates.append({
                    "span_text":             result_str,
                    "replacement":           wrong_str,
                    "subtype":               "DIFFERENZ_INLINE",
                    "is_reformulation":      False,
                    "requires_human_review": False,
                    "rationale": (
                        f"Korrekte Differenz {a_str}–{b_str}={_format_num(correct, result_str)}; "
                        f"injiziert: {wrong_str}"
                    ),
                })
        except (ValueError, ZeroDivisionError):
            continue

    # Muster 3: "A × B = C"
    for m in _CALC_MUL_RE.finditer(segment_text):
        a_str, b_str, result_str = m.group(1), m.group(2), m.group(3)
        try:
            a, b, result = _parse_num(a_str), _parse_num(b_str), _parse_num(result_str)
            correct = a * b
            if abs(result - correct) < max(0.1, abs(correct) * 0.02):
                delta = max(1.0, correct * 0.1)
                wrong = correct + delta
                wrong_str = _format_num(wrong, result_str)
                candidates.append({
                    "span_text":             result_str,
                    "replacement":           wrong_str,
                    "subtype":               "PRODUKT",
                    "is_reformulation":      False,
                    "requires_human_review": False,
                    "rationale": (
                        f"Korrektes Produkt {a_str}×{b_str}={_format_num(correct, result_str)}; "
                        f"injiziert: {wrong_str}"
                    ),
                })
        except (ValueError, ZeroDivisionError):
            continue

    # Muster 4: "d m / t s × 3.6 = V km/h" — Geschwindigkeit aus Weg und Zeit
    for m in _CALC_SPEED_DIV_RE.finditer(segment_text):
        d_str, t_str, v_str, unit = m.group(1), m.group(2), m.group(3), m.group(4)
        try:
            d = _parse_num(d_str)
            t = _parse_num(t_str)
            v_doc = _parse_num(v_str)
            if t == 0:
                continue
            if unit.lower() == "km/h":
                correct = _trunc1(d / t * 3.6)
            else:  # m/s
                correct = _trunc1(d / t)
            # Dokument korrekt (±2% Toleranz) → Fehler injizieren
            if abs(v_doc - correct) <= max(0.2, abs(correct) * 0.02):
                # Falschen Wert: ±10% abweichend, aber plausibel
                delta = max(0.5, abs(correct) * 0.10)
                wrong = _trunc1(correct + delta)
                wrong_str = _format_num(wrong, v_str)
                candidates.append({
                    "span_text":             v_str,
                    "replacement":           wrong_str,
                    "subtype":               "GESCHWINDIGKEIT",
                    "is_reformulation":      False,
                    "requires_human_review": False,
                    "rationale": (
                        f"Korrekte Geschwindigkeit {d_str}m/{t_str}s×3.6"
                        f"={correct}{unit}; injiziert: {wrong_str}{unit}"
                    ),
                })
        except (ValueError, ZeroDivisionError):
            continue

    # Muster 5: allgemeine Division "A / B = C" (z.B. Weg/Zeit, Verhältnisse)
    for m in _CALC_DIV_RE.finditer(segment_text):
        a_str, b_str, result_str = m.group(1), m.group(2), m.group(3)
        # Überlapp mit anderen Mustern verhindern
        span_start = m.start()
        already_covered = any(
            ex.start() <= span_start <= ex.end()
            for ex in list(_CALC_DIFF_RE.finditer(segment_text))
            + list(_CALC_SIMPLE_RE.finditer(segment_text))
            + list(_CALC_MUL_RE.finditer(segment_text))
        )
        if already_covered:
            continue
        try:
            a, b, result = _parse_num(a_str), _parse_num(b_str), _parse_num(result_str)
            if b == 0:
                continue
            correct = _trunc1(a / b)
            if abs(result - correct) <= max(0.1, abs(correct) * 0.02):
                delta = max(0.5, abs(correct) * 0.10)
                wrong = _trunc1(correct + delta)
                wrong_str = _format_num(wrong, result_str)
                candidates.append({
                    "span_text":             result_str,
                    "replacement":           wrong_str,
                    "subtype":               "DIVISION",
                    "is_reformulation":      False,
                    "requires_human_review": False,
                    "rationale": (
                        f"Korrekte Division {a_str}/{b_str}"
                        f"={_format_num(correct, result_str)}; injiziert: {wrong_str}"
                    ),
                })
        except (ValueError, ZeroDivisionError):
            continue

    return candidates


# ── RECHT_ABSICHERUNG  (Regex-Filter + LLM) ───────────────────────────────────

_HEDGING_RE = re.compile(
    r"\b(?:vermutlich|möglicherweise|wahrscheinlich|könnte|dürfte|scheint|"
    r"lässt\s+vermuten|deutet\s+darauf|spricht\s+dafür|gemäss|laut\b|"
    r"zufolge|nach\s+Angaben|unserer\s+Einschätzung|gehen\s+davon\s+aus|"
    r"geht\s+davon\s+aus|erachten\s+als|beurteilen\s+als)\b",
    re.IGNORECASE,
)


def _split_sentences(text: str) -> List[str]:
    return [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if len(s.strip()) > 20]


def extract_candidates_absicherung(segment_text: str) -> List[Dict]:
    """
    Findet Sätze mit Hedging-Sprache, die zu absolut umformuliert werden sollen.
    Satz-Level-Check: der Kandidatensatz muss selbst ein Hedging-Wort enthalten
    (nicht nur irgendein Satz im Segment).
    Zusätzlich: Sätze die bereits im Indikativ eine absolute Tatsachenbehauptung
    sind, werden ausgeschlossen (kein Hedging zum Entfernen vorhanden).
    """
    candidates = []
    for sent in _split_sentences(segment_text):
        # Satz muss selbst Hedging enthalten
        if not _HEDGING_RE.search(sent):
            continue
        # Satz muss lang genug sein um sinnvoll umformuliert werden zu können
        if len(sent.strip()) < 20:
            continue
        candidates.append({
            "span_text":             sent,
            "replacement":           "",
            "subtype":               "HEDGING",
            "is_reformulation":      True,
            "requires_human_review": False,
            "rationale":             "",
        })
    return candidates


# ── RECHT_TRENNUNG  (Regex-Filter + LLM) ──────────────────────────────────────

_KONJUNKTIV_RE = re.compile(
    r"\b(?:hätte|wäre|könnte|würde|sollte|müsste|dürfte|"
    r"hätten|wären|könnten|würden|sollten|müssten)\b",
    re.IGNORECASE,
)
_BEWERTUNG_RE = re.compile(
    r"\b(?:Bewertung|Einschätzung|Beurteilung|Hypothese|Annahme|"
    r"Schlussfolgerung|erachten|beurteilen|erscheint|wirkt)\b",
    re.IGNORECASE,
)


def extract_candidates_trennung(segment_text: str) -> List[Dict]:
    """
    Findet Sätze mit Konjunktiv II oder Bewertungssprache,
    die als Tatsache umformuliert werden sollen.

    Satz-Level-Check: der Kandidatensatz muss selbst den Konjunktiv oder die
    Bewertungssprache enthalten – nicht nur irgendein anderer Satz im Segment.
    Das verhindert dass Indikativ-Sätze als Kandidaten gewählt werden, weil
    ein anderer Satz im gleichen Segment einen Konjunktiv enthält.
    """
    candidates = []
    for sent in _split_sentences(segment_text):
        sent_stripped = sent.strip()
        if len(sent_stripped) < 20:
            continue

        # KONJUNKTIV: Satz muss selbst einen Konjunktiv-II-Ausdruck enthalten
        if _KONJUNKTIV_RE.search(sent_stripped):
            candidates.append({
                "span_text":             sent_stripped,
                "replacement":           "",
                "subtype":               "KONJUNKTIV",
                "is_reformulation":      True,
                "requires_human_review": False,
                "rationale":             "",
            })
        # BEWERTUNG: Satz muss sowohl Bewertungswort als auch Hedging enthalten
        elif _BEWERTUNG_RE.search(sent_stripped) and _HEDGING_RE.search(sent_stripped):
            candidates.append({
                "span_text":             sent_stripped,
                "replacement":           "",
                "subtype":               "BEWERTUNG",
                "is_reformulation":      True,
                "requires_human_review": False,
                "rationale":             "",
            })
    return candidates


# ── HYPO_INKONSISTENZ  (LLM, braucht ganzen Text) ─────────────────────────────

_HYPO_MARKER_RE = re.compile(
    r"\b(?:H1|H2|H3|Hypothese|Identitätshypothese|Nichtidentitätshypothese|"
    r"stützt|widerspricht|sehr\s+stark|klar\s+gestützt|spricht\s+(?:für|gegen))\b",
    re.IGNORECASE,
)


def extract_candidates_hypo(segment_text: str) -> List[Dict]:
    """Findet Segmente mit Hypothesenbewertungen."""
    if not _HYPO_MARKER_RE.search(segment_text):
        return []
    # Ganzes Segment als Kandidat – LLM soll Bewertung kippen
    return [{
        "span_text":             segment_text,
        "replacement":           "",
        "subtype":               "HYPO",
        "is_reformulation":      True,
        "requires_human_review": True,    # zwingend manuell prüfen
        "rationale":             "",
    }]


# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 2 – BUDGET-SELEKTOR
# ═══════════════════════════════════════════════════════════════════════════════

def select_class(
        available_classes: List[str],
        produced_counts: Dict[str, int],
        total_produced: int,
        allowed_classes: Optional[List[str]] = None,
) -> Optional[str]:
    """
    Wählt die Fehlerklasse mit dem grössten Budget-Defizit
    (Zielanteil − Istanteil), beschränkt auf available_classes ∩ allowed_classes.
    Hard Cap: kein Typ darf mehr als 2× seinen Zielanteil überschreiten.
    """
    permitted = set(allowed_classes) if allowed_classes else set(CLASS_PRIORITY)
    eligible = [c for c in available_classes if c in permitted]
    if not eligible:
        return None

    denom = max(total_produced, 1)
    scores: Dict[str, float] = {}
    for cls in eligible:
        target = CLASS_BUDGET.get(cls, 0.0)
        actual = produced_counts.get(cls, 0) / denom
        if actual > target * 2.0:
            continue
        prio_bonus = (
            (len(CLASS_PRIORITY) - CLASS_PRIORITY.index(cls)) * 0.001
            if cls in CLASS_PRIORITY else 0
        )
        scores[cls] = (target - actual) + prio_bonus

    return max(scores, key=lambda c: scores[c]) if scores else None


# ═══════════════════════════════════════════════════════════════════════════════
# PHASE 3 – INJEKTION  (pro Fehlerklasse)
# ═══════════════════════════════════════════════════════════════════════════════

# ── LLM-Client ────────────────────────────────────────────────────────────────

def make_injection_client() -> OllamaClient:
    base_url = require_env("OLLAMA_BASE_URL")
    model    = require_env("INJECTION_MODEL")
    timeout  = env_int("LLM_TIMEOUT_S", 300)
    temp     = float(env_str("INJECTION_TEMPERATURE", "0.3"))
    options  = {"temperature": temp}
    options.update(env_json_object_optional("LLM_OPTIONS_JSON"))
    return OllamaClient(base_url=base_url, model=model, options=options, timeout_s=timeout)


# ── LLM-Debug-Logging ─────────────────────────────────────────────────────────
# Aktivieren: LLM_DEBUG=true in .env oder als Umgebungsvariable.
# Gibt System-Prompt, User-Prompt und Raw-Response für jeden LLM-Call aus.

_LLM_CALL_COUNTER = 0   # fortlaufende Nummer pro Lauf

def _llm_call(
        llm: OllamaClient,
        messages: List[Dict[str, str]],
        label: str = "",
        schema: Optional[Dict[str, Any]] = None,
) -> str:
    """
    Zentraler Wrapper für alle LLM-Aufrufe.
    Loggt Input und Output wenn LLM_DEBUG=true gesetzt ist.
    label: kurzer Bezeichner (z.B. "STRUKT Phase2", "Plausibilitätsprüfung")
    """
    global _LLM_CALL_COUNTER
    _LLM_CALL_COUNTER += 1
    call_nr = _LLM_CALL_COUNTER
    debug = env_bool("LLM_DEBUG", False)

    if debug:
        sep = "─" * 64
        print(f"\n┌{sep}")
        print(f"│ LLM-Call #{call_nr}  [{label}]")
        print(f"├{sep}")
        for msg in messages:
            role = msg["role"].upper()
            content = msg["content"]
            print(f"│ ── {role} ──────────────────────────────────────────────")
            for line in content.splitlines():
                print(f"│   {line}")
        print(f"├{sep}")

    raw = llm.chat(messages, schema=schema)

    if debug:
        print(f"│ ── RESPONSE ─────────────────────────────────────────────")
        for line in (raw or "").splitlines():
            print(f"│   {line}")
        print(f"└{sep}\n")

    return raw


# ── JSON-Schema für grammar-constrained Decoding ──────────────────────────────

def _base_schema() -> Dict[str, Any]:
    return {
        "type": "object",
        "properties": {
            "has_candidate": {"type": "boolean"},
            "original_span": {"type": "string"},
            "injected_span": {"type": "string"},
            "severity_id":   {"type": "string", "enum": ["HIGH", "MEDIUM", "LOW"]},
            "rationale":     {"type": "string"},
        },
        "required": ["has_candidate", "original_span", "injected_span",
                     "severity_id", "rationale"],
    }

def _phase1_schema() -> Dict[str, Any]:
    """JSON-Schema für Phase-1-Call: Kandidatensatz identifizieren."""
    return {
        "type": "object",
        "properties": {
            "found":      {"type": "boolean"},
            "candidate":  {"type": "string"},
        },
        "required": ["found", "candidate"],
    }




def _parse_llm_json(raw: str) -> Optional[Dict[str, Any]]:
    content = raw.strip()
    if "```" in content:
        for part in content.split("```"):
            p = part.lstrip("json").strip()
            if p.startswith("{"):
                content = p
                break
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        return None


# ── Deterministisch: FORMALES_REDAKTION ───────────────────────────────────────

def inject_formales_llm(
        segment_text: str,
        llm: OllamaClient,
        case_id: str,
        seg_idx: int,
        proposal_counter: int,
        rng: random.Random,
        eligible_subtypes: Optional[List[str]] = None,
        formales_candidates_meta: Optional[Dict[str, Any]] = None,
) -> Optional[Dict]:
    """
    LLM-basierte Injektion für FORMALES_REDAKTION.
    Subtyp aus Precheck-validierten eligible_subtypes; für DOPPELKONSONANT
    werden zusätzlich die konkreten Kandidatenwörter in den Prompt eingebettet.
    """
    pool        = eligible_subtypes if eligible_subtypes else _FORMALES_SUBTYPES
    subtype     = rng.choice(pool)
    description = _FORMALES_SUBTYPE_DESC[subtype]
    # Für DOPPELKONSONANT: Kandidatenwörter aus dem Precheck einbetten
    candidate_words = formales_candidates_meta.get(subtype, {}).get("candidate_words", []) \
        if formales_candidates_meta else []

    system = (
        "Du bist ein Assistent, der synthetische Rechtschreibfehler in deutsche "
        "forensische Gutachtentexte injiziert.\n\n"
        "## Klasse: FORMALES_REDAKTION\n"
        "Sprachliche Fehler ohne inhaltliche Auswirkung.\n\n"
        "## Regeln\n"
        "1. Ändere AUSSCHLIESSLICH einen einzigen Span (ein Wort oder eine kurze Wortgruppe).\n"
        "2. Der Fehler darf den Inhalt nicht verändern – nur die Form.\n"
        "3. Fachbegriffe, Eigennamen, Abkürzungen und Zahlen nicht verändern.\n"
        "4. Kein Kandidat vorhanden → has_candidate=false, alle anderen Felder leer.\n"
        "5. Antworte NUR mit dem JSON-Objekt."
    )
    # Für alle Subtypen mit Wort-Extraktor: Kandidatenwörter explizit nennen.
    # Das Modell darf NUR aus dieser Liste wählen – verhindert Halluzination
    # (z.B. Umlaut-Fehler auf Wörter ohne Umlaut, Tippfehler auf Eigennamen).
    _SUBTYPE_HINT_LABELS = {
        "DOPPELKONSONANT": "Wörter mit Doppelkonsonant",
        "TIPPFEHLER":      "Geeignete Kleinwörter für Tippfehler",
        "ADJEKTIV_GROSS":  "Mögliche Adjektive (Flexionsendung -en/-er/-em)",
        "HOMOPHON":        "Homophon-Kandidaten im Segment",
    }
    word_hint = ""
    if candidate_words and subtype in _SUBTYPE_HINT_LABELS:
        label = _SUBTYPE_HINT_LABELS[subtype]
        word_hint = (
                f"{label} in diesem Segment: "
                + ", ".join(f'"{w}"' for w in candidate_words)
                + "\nWähle den original_span AUSSCHLIESSLICH aus dieser Liste.\n"
                + "Eigennamen (Personen, Ortsnamen) NIEMALS verändern.\n\n"
        )
    user = (
        f"Geforderter Fehlersubtyp: {subtype}\n"
        f"{word_hint}"
        f"Anweisung: {description}\n\n"
        f"SEGMENT:\n\"\"\"\n{segment_text.strip()}\n\"\"\"\n\n"
        f"Injiziere einen {subtype}-Fehler. "
        "original_span = korrekter Originaltext, injected_span = fehlerhafter Text."
    )

    try:
        raw = _llm_call(
            llm,
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            label="FORMALES Injektion",
            schema=_base_schema(),
        )
    except Exception as e:
        print(f"    [FEHLER] LLM: {e}")
        return None

    result = _parse_llm_json(raw)
    if not result or not result.get("has_candidate"):
        return None

    original_span = result.get("original_span", "").strip()
    injected_span = result.get("injected_span", "").strip()

    if not original_span or original_span not in segment_text:
        return None

    meta = CLASS_META["FORMALES_REDAKTION"]
    return _make_proposal(
        case_id=case_id, seg_idx=seg_idx, counter=proposal_counter,
        subclass_id="FORMALES_REDAKTION",
        segment_text=segment_text,
        original_span=original_span,
        injected_span=injected_span,
        severity_id=meta["severity_id"],
        change_type_id=meta["change_type_id"],
        rationale=result.get("rationale", f"{subtype}: '{original_span}' → '{injected_span}'"),
        subtype=subtype,
        requires_human_review=False,
    )


# ── Deterministisch: RECHEN_ARITHMETIK ────────────────────────────────────────

def inject_rechen(
        candidate: Dict,
        segment_text: str,
        case_id: str,
        seg_idx: int,
        proposal_counter: int,
) -> Optional[Dict]:
    span  = candidate["span_text"]
    wrong = candidate["replacement"]
    if not span or not wrong or span not in segment_text:
        return None
    meta = CLASS_META["RECHEN_ARITHMETIK"]
    return _make_proposal(
        case_id=case_id, seg_idx=seg_idx, counter=proposal_counter,
        subclass_id="RECHEN_ARITHMETIK",
        segment_text=segment_text,
        original_span=span,
        injected_span=wrong,
        severity_id=meta["severity_id"],
        change_type_id=meta["change_type_id"],
        rationale=candidate["rationale"],
        subtype=candidate["subtype"],
        requires_human_review=False,
    )


# ── LLM: STRUKT_BEFUND_BESCHREIBUNG ───────────────────────────────────────────

_STRUKT_ENTITY_INSTRUCTIONS: Dict[str, str] = {
    "MESSWERT": (
        "Suche einen Zahlenwert MIT physikalischer Einheit (Hz, kHz, m, kg, dB, Promille, Bit etc.).\n"
        "Ändere NUR die Einheit.\n"
        "Beispiele: '44.1 kHz' → '45.1 Hz', '10 mm' → '10 dm', '3 Stunden' → '3 Minuten'.\n"
        "Der Fehler soll wie ein realistischer Tippfehler beim Abtippen wirken."
    ),
    "INSTITUTION": (
        "Suche eine benannte Institution. Ersetze mit einer Institution aus der Whitelist,\n"
        "die klar eine ANDERE Organisation ist (anderer Kanton oder Zweck).\n"
        "Zu ähnliche Namen sind abgelehnt."
    ),
    "PERSON": (
        "Die bekannten Personen sind im User-Prompt aufgeführt.\n"
        "Ersetze genau einen bekannten Namen durch einen fiktiven anderen Namen.\n"
        "\n"
        "REGELN:\n"
        "  - Ändere den Vornamen zu einem anderen echten deutschen Vornamen.\n"
        "  - Der neue Vorname muss sich mindestens 4 Buchstaben vom Original unterscheiden.\n"
        "  - Nachname und Titel bleiben unverändert.\n"
        "  - Der neue Name darf NICHT bereits im Dokument vorkommen —\n"
        "    weder als vollständiger Name noch als Teilstring oder in anderer Reihenfolge.\n"
        "  - Der neue Name darf NICHT aus Teilen anderer im Dokument genannter Personen\n"
        "    zusammengesetzt sein (z.B. Vorname von Person A + Nachname von Person B).\n"
        "\n"
        "VERBOTEN:\n"
        "  - Roland → Rolf (zu ähnlich, < 4 Buchstaben Unterschied)\n"
        "  - Roland → Roli (Kurzform desselben Namens)\n"
        "  - Verwendung von Namen oder Namensteilen die bereits im Dokument stehen.\n"
        "\n"
        "ERLAUBT (Beispiele):\n"
        "  - Roland → Markus, Stefan → Thomas, Marie → Anna, Peter → Heinrich"
    ),
    "DATUM": (
        "Suche ein Datum mit ausgeschriebenem Monat.\n"
        "Ändere NUR die Jahreszahl zu einem klar falschen Jahr\n"
        "(Zukunft nach 2026 ODER >5 Jahre zu früh). Tag und Monat nie ändern."
    ),
}


def inject_strukt_llm(
        subtype: str,
        candidate_span: str,
        segment_text: str,
        entity_whitelist: List[str],
        reference_facts: Optional[Dict],
        llm: OllamaClient,
        case_id: str,
        seg_idx: int,
        proposal_counter: int,
) -> Optional[Dict]:
    """
    Phase 3 für STRUKT_BEFUND_BESCHREIBUNG.
    candidate_span ist der von Phase 1 extrahierte Originalspan – das LLM
    muss ihn nur noch ersetzen, nicht mehr selbst suchen.
    """
    whitelist_block = "\n".join(f"  - {e}" for e in entity_whitelist)
    system = (
        "Du bist ein Assistent, der synthetische Fehler in forensische Gutachtentexte injiziert.\n\n"
        "## Klasse: STRUKT_BEFUND_BESCHREIBUNG\n"
        "Ein Sachverhalt wird falsch beschrieben: falscher Messwert, falsche Institution,\n"
        "falsche Person oder falsches Datum.\n\n"
        "## Regeln\n"
        "1. Ersetze AUSSCHLIESSLICH den im User-Prompt angegebenen Span.\n"
        "   Der original_span im JSON MUSS exakt diesem Span entsprechen.\n"
        "2. Der injizierte Fehler muss einem Referenzfakt widersprechen ODER physikalisch unmöglich sein.\n"
        f"3. Bei INSTITUTION: Ersatz NUR aus dieser Whitelist:\n{whitelist_block}\n"
        "4. Bei PERSON: Ersatz NUR unter Variation des im User-Prompt genannten Namens.\n"
        "5. Kein sinnvoller Ersatz möglich → has_candidate=false, alle anderen Felder leer.\n"
        "6. Antworte NUR mit dem JSON-Objekt."
    )

    instruction = _STRUKT_ENTITY_INSTRUCTIONS.get(subtype, "Substituiere den angegebenen Span.")
    user = (
        f"Zu ersetzender Span: \"{candidate_span}\"\n"
        f"Entitätstyp: {subtype}\n"
        f"Anweisung: {instruction}\n\n"
        f"SEGMENT (Kontext):\n\"\"\"\n{segment_text.strip()}\n\"\"\"\n\n"
        f"Ersetze ausschliesslich \"{candidate_span}\" durch einen falschen Wert. "
        f"Setze original_span=\"{candidate_span}\" und liefere den Ersatz in injected_span."
    )

    try:
        raw = _llm_call(
            llm,
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            label="STRUKT Injektion",
            schema=_base_schema(),
        )
    except Exception as e:
        print(f"    [FEHLER] LLM: {e}")
        return None

    result = _parse_llm_json(raw)
    if not result or not result.get("has_candidate"):
        return None

    meta = CLASS_META["STRUKT_BEFUND_BESCHREIBUNG"]
    return _make_proposal(
        case_id=case_id, seg_idx=seg_idx, counter=proposal_counter,
        subclass_id="STRUKT_BEFUND_BESCHREIBUNG",
        segment_text=segment_text,
        original_span=result.get("original_span", "").strip(),
        injected_span=result.get("injected_span", "").strip(),
        severity_id=result.get("severity_id", meta["severity_id"]),
        change_type_id=meta["change_type_id"],
        rationale=result.get("rationale", "").strip(),
        subtype=subtype,
        requires_human_review=False,
    )


# ── LLM: RECHT_ABSICHERUNG ────────────────────────────────────────────────────

def inject_absicherung_llm(
        candidate: Dict,
        segment_text: str,
        llm: OllamaClient,
        case_id: str,
        seg_idx: int,
        proposal_counter: int,
) -> Optional[Dict]:
    """
    Zweiphasen-Injektion für RECHT_ABSICHERUNG:
    Phase 1 – LLM wählt den besten Kandidatensatz aus dem Segment.
    Phase 2 – LLM schreibt NUR diesen Satz zu absolut um.
    """
    # ── Phase 1: Kandidatensatz identifizieren ────────────────────────────
    p1_system = (
        "Du analysierst forensische Gutachtentexte.\n"
        "Finde den Satz im Segment, der am ehesten für einen RECHT_ABSICHERUNG-Fehler geeignet ist:\n"
        "Sätze mit Hedging-Sprache (vermutlich, könnte, gemäss, zufolge, erscheine, seien etc.).\n"
        "Antworte NUR mit JSON: {found: bool, candidate: string}\n"
        "candidate = exakter Originaltext des Satzes, leer wenn keiner gefunden."
    )
    p1_user = f"SEGMENT:\n\"\"\"{segment_text.strip()}\"\"\"\n\nWelcher Satz eignet sich für RECHT_ABSICHERUNG?"
    try:
        p1_raw = _llm_call(
            llm,
            [{"role": "system", "content": p1_system}, {"role": "user", "content": p1_user}],
            label="RECHT_ABSICHERUNG Phase 1",
            schema=_phase1_schema(),
        )
    except Exception as e:
        print(f"    [FEHLER] Phase 1 LLM: {e}")
        return None
    p1 = _parse_llm_json(p1_raw)
    if not p1 or not p1.get("found") or not p1.get("candidate", "").strip():
        return None
    chosen = p1["candidate"].strip()
    if chosen not in segment_text:
        chosen = candidate["span_text"]   # Fallback auf Regex-Kandidat

    # ── Phase 2: Gewählten Satz zu absolut umschreiben ────────────────────
    p2_system = (
        "Du bist ein Assistent, der synthetische Fehler der Klasse RECHT_ABSICHERUNG erzeugt.\n\n"
        "## Definition\n"
        "Eine Aussage wird zu absolut formuliert: Hedging, Konjunktiv und Relativierungen\n"
        "werden entfernt, so dass eine Tatsachenbehauptung entsteht.\n\n"
        "## Regeln\n"
        "1. Schreibe AUSSCHLIESSLICH den angegebenen Satz um.\n"
        "2. Entferne alle Hedging-Wörter (vermutlich, könnte, gemäss, zufolge etc.).\n"
        "3. Formuliere als gesicherte Tatsache.\n"
        "4. Der umgeschriebene Satz muss im Kontext unangemessen absolut wirken.\n"
        "5. Antworte NUR mit dem JSON-Objekt."
    )
    p2_user = (
        f"Satz zum Umschreiben:\n\"\"\"\n{chosen}\n\"\"\"\n\n"
        "Schreibe den Satz zu absolut um. "
        f"Setze original_span exakt auf: \"{chosen}\""
    )
    try:
        raw = _llm_call(
            llm,
            [{"role": "system", "content": p2_system}, {"role": "user", "content": p2_user}],
            label="RECHT_ABSICHERUNG Phase 2",
            schema=_base_schema(),
        )
    except Exception as e:
        print(f"    [FEHLER] Phase 2 LLM: {e}")
        return None

    result = _parse_llm_json(raw)
    if not result or not result.get("has_candidate"):
        return None

    meta = CLASS_META["RECHT_ABSICHERUNG"]
    return _make_proposal(
        case_id=case_id, seg_idx=seg_idx, counter=proposal_counter,
        subclass_id="RECHT_ABSICHERUNG",
        segment_text=segment_text,
        original_span=chosen,
        injected_span=result.get("injected_span", "").strip(),
        severity_id=meta["severity_id"],
        change_type_id=meta["change_type_id"],
        rationale=result.get("rationale", "Hedging entfernt – Aussage zu absolut formuliert."),
        subtype="ABSOLUT",
        requires_human_review=False,
    )


# ── LLM: RECHT_TRENNUNG ───────────────────────────────────────────────────────

def inject_trennung_llm(
        candidate: Dict,
        segment_text: str,
        llm: OllamaClient,
        case_id: str,
        seg_idx: int,
        proposal_counter: int,
) -> Optional[Dict]:
    """
    Zweiphasen-Injektion für RECHT_TRENNUNG:
    Phase 1 – LLM wählt den besten Konjunktiv/Bewertungs-Satz.
    Phase 2 – LLM wandelt ihn in eine Tatsachenaussage um.
    """
    # ── Phase 1: Kandidatensatz identifizieren ────────────────────────────
    p1_system = (
        "Du analysierst forensische Gutachtentexte.\n"
        "Finde den Satz im Segment, der am ehesten für RECHT_TRENNUNG geeignet ist:\n"
        "Sätze mit Konjunktiv II (hätte, wäre, würde, könnte) oder epistemischen Markern\n"
        "('unserer Ansicht nach', 'erachten wir', 'erscheint', 'aus unserer Sicht').\n"
        "Antworte NUR mit JSON: {found: bool, candidate: string}\n"
        "candidate = exakter Originaltext des Satzes, leer wenn keiner gefunden."
    )
    p1_user = f"SEGMENT:\n\"\"\"{segment_text.strip()}\"\"\"\n\nWelcher Satz eignet sich für RECHT_TRENNUNG?"
    try:
        p1_raw = _llm_call(
            llm,
            [{"role": "system", "content": p1_system}, {"role": "user", "content": p1_user}],
            label="RECHT_TRENNUNG Phase 1",
            schema=_phase1_schema(),
        )
    except Exception as e:
        print(f"    [FEHLER] Phase 1 LLM: {e}")
        return None
    p1 = _parse_llm_json(p1_raw)
    if not p1 or not p1.get("found") or not p1.get("candidate", "").strip():
        return None
    chosen = p1["candidate"].strip()
    if chosen not in segment_text:
        chosen = candidate["span_text"]   # Fallback auf Regex-Kandidat

    # ── Phase 2: Gewählten Satz in Indikativ umschreiben ─────────────────
    p2_system = (
        "Du bist ein Assistent, der synthetische Fehler der Klasse RECHT_TRENNUNG erzeugt.\n\n"
        "## Definition\n"
        "Eine Bewertung oder Hypothese wird als gesicherte Tatsache dargestellt.\n\n"
        "## Regeln\n"
        "1. Schreibe AUSSCHLIESSLICH den angegebenen Satz um.\n"
        "2. Wandle Konjunktiv II in Indikativ um.\n"
        "3. Entferne epistemische Marker.\n"
        "4. Antworte NUR mit dem JSON-Objekt."
    )
    p2_user = (
        f"Satz zum Umschreiben:\n\"\"\"\n{chosen}\n\"\"\"\n\n"
        "Schreibe Konjunktiv → Indikativ, Bewertung → Tatsache. "
        f"Setze original_span exakt auf: \"{chosen}\""
    )
    try:
        raw = _llm_call(
            llm,
            [{"role": "system", "content": p2_system}, {"role": "user", "content": p2_user}],
            label="RECHT_TRENNUNG Phase 2",
            schema=_base_schema(),
        )
    except Exception as e:
        print(f"    [FEHLER] Phase 2 LLM: {e}")
        return None

    result = _parse_llm_json(raw)
    if not result or not result.get("has_candidate"):
        return None

    meta = CLASS_META["RECHT_TRENNUNG"]
    return _make_proposal(
        case_id=case_id, seg_idx=seg_idx, counter=proposal_counter,
        subclass_id="RECHT_TRENNUNG",
        segment_text=segment_text,
        original_span=chosen,
        injected_span=result.get("injected_span", "").strip(),
        severity_id=meta["severity_id"],
        change_type_id=meta["change_type_id"],
        rationale=result.get("rationale",
                             "Bewertung als Tatsache dargestellt (Konjunktiv → Indikativ)."),
        subtype=candidate["subtype"],
        requires_human_review=False,
    )


# ── LLM: HYPO_INKONSISTENZ ────────────────────────────────────────────────────

def inject_hypo_llm(
        candidate: Dict,
        segment_text: str,
        llm: OllamaClient,
        case_id: str,
        seg_idx: int,
        proposal_counter: int,
) -> Optional[Dict]:
    """
    Zweiphasen-Injektion für HYPO_INKONSISTENZ:
    Phase 1 – LLM identifiziert den konkreten Bewertungsausdruck.
    Phase 2 – LLM kippt nur diesen Ausdruck.
    """
    # ── Phase 1: Bewertungsausdruck identifizieren ────────────────────────
    p1_system = (
        "Du analysierst forensische Gutachtentexte.\n"
        "Finde den Bewertungsausdruck einer Hypothese im Segment.\n"
        "Typisch: 'spricht stark für', 'stützt', 'widerspricht', 'klar gestützt', 'leicht für/gegen'.\n"
        "Antworte NUR mit JSON: {found: bool, candidate: string}\n"
        "candidate = exakter Originaltext des Bewertungsausdrucks (möglichst kurz, max. 1 Satz)."
    )
    p1_user = f"SEGMENT:\n\"\"\"{segment_text.strip()}\"\"\"\n\nWelcher Bewertungsausdruck soll gekippt werden?"
    try:
        p1_raw = _llm_call(
            llm,
            [{"role": "system", "content": p1_system}, {"role": "user", "content": p1_user}],
            label="HYPO_INKONSISTENZ Phase 1",
            schema=_phase1_schema(),
        )
    except Exception as e:
        print(f"    [FEHLER] Phase 1 LLM: {e}")
        return None
    p1 = _parse_llm_json(p1_raw)
    if not p1 or not p1.get("found") or not p1.get("candidate", "").strip():
        return None
    chosen = p1["candidate"].strip()
    if chosen not in segment_text:
        return None   # kein Fallback bei HYPO – zu riskant

    # ── Phase 2: Bewertungsausdruck kippen ────────────────────────────────
    p2_system = (
        "Du bist ein Assistent, der synthetische Fehler der Klasse HYPO_INKONSISTENZ erzeugt.\n\n"
        "## Regeln\n"
        "1. Ersetze NUR den angegebenen Span.\n"
        "2. Kehre die Bewertung um: 'für' → 'gegen', 'gestützt' → 'widersprochen' etc.\n"
        "3. Ändere NICHTS ausserhalb des Spans.\n"
        "4. Antworte NUR mit dem JSON-Objekt."
    )
    p2_user = (
        f"Zu kippender Span:\n\"\"\"\n{chosen}\n\"\"\"\n\n"
        "Kehre die Hypothesenbewertung um. "
        f"Setze original_span exakt auf: \"{chosen}\""
    )
    try:
        raw = _llm_call(
            llm,
            [{"role": "system", "content": p2_system}, {"role": "user", "content": p2_user}],
            label="HYPO_INKONSISTENZ Phase 2",
            schema=_base_schema(),
        )
    except Exception as e:
        print(f"    [FEHLER] Phase 2 LLM: {e}")
        return None

    result = _parse_llm_json(raw)
    if not result or not result.get("has_candidate"):
        return None

    meta = CLASS_META["HYPO_INKONSISTENZ"]
    return _make_proposal(
        case_id=case_id, seg_idx=seg_idx, counter=proposal_counter,
        subclass_id="HYPO_INKONSISTENZ",
        segment_text=segment_text,
        original_span=chosen,
        injected_span=result.get("injected_span", "").strip(),
        severity_id=meta["severity_id"],
        change_type_id=meta["change_type_id"],
        rationale=result.get("rationale", "Hypothesenbewertung gekippt."),
        subtype="HYPO",
        requires_human_review=True,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# PROPOSAL-HILFSFUNKTIONEN
# ═══════════════════════════════════════════════════════════════════════════════

def _make_proposal(
        case_id: str,
        seg_idx: int,
        counter: int,
        subclass_id: str,
        segment_text: str,
        original_span: str,
        injected_span: str,
        severity_id: str,
        change_type_id: str,
        rationale: str,
        subtype: str,
        requires_human_review: bool,
) -> Dict[str, Any]:
    proposal_id = f"{case_id}_seg_{seg_idx:04d}_inj_{counter:02d}"
    meta = CLASS_META.get(subclass_id, {})
    proposal = {
        "proposal_id":           proposal_id,
        "case_id":               case_id,
        "segment_index":         seg_idx,
        "segment_text":          segment_text,
        "subclass_id":           subclass_id,
        "change_type_id":        change_type_id or meta.get("change_type_id", ""),
        "original_span":         original_span,
        "injected_span":         injected_span,
        "severity_id":           severity_id or meta.get("severity_id", "MEDIUM"),
        "rationale":             rationale,
        "subtype":               subtype,
        "requires_human_review": requires_human_review,
        "status":                "accepted",
        "validation_ok":         None,
        "validation_issues":     [],
    }
    return proposal


def _known_person_names(reference_facts: Optional[Dict[str, Any]]) -> List[str]:
    if not reference_facts:
        return []
    personen = reference_facts.get("facts", {}).get("personen", [])
    return [
        p["name"].strip() for p in personen
        if p.get("confidence") in {"high", "medium"} and p.get("name", "").strip()
    ]


# ═══════════════════════════════════════════════════════════════════════════════
# VALIDIERUNG
# ═══════════════════════════════════════════════════════════════════════════════

def validate_proposal(
        proposal: Dict[str, Any],
        segment_text: str,
        reference_facts: Optional[Dict[str, Any]] = None,
) -> List[str]:
    issues: List[str] = []
    original_span = (proposal.get("original_span") or "").strip()
    injected_span = (proposal.get("injected_span") or "").strip()
    subclass_id   = proposal.get("subclass_id", "")
    subtype       = proposal.get("subtype", "")

    if not original_span:
        issues.append("original_span ist leer")
    elif original_span not in segment_text:
        issues.append(f"original_span nicht im Segment: '{original_span[:80]}'")

    if not injected_span:
        issues.append("injected_span ist leer")

    if original_span and injected_span and original_span.strip() == injected_span.strip():
        issues.append("original_span und injected_span identisch – kein Fehler injiziert")
    # Auch bei Reformulierungen: injected darf nicht leer oder nur Whitespace sein
    if injected_span and not injected_span.strip():
        issues.append("injected_span ist nur Whitespace")

    if not proposal.get("rationale", "").strip():
        issues.append("rationale fehlt")

    # PERSON: Referenzfakten-Check
    if subtype == "PERSON" or subclass_id == "STRUKT_BEFUND_BESCHREIBUNG" and subtype == "PERSON":
        known = _known_person_names(reference_facts)
        if not known:
            issues.append("PERSON-Injektion ohne Referenzfakten")
        else:
            orig_l = original_span.lower()
            if not any(n.lower() in orig_l or orig_l in n.lower() for n in known):
                issues.append(
                    f"original_span '{original_span}' ist kein Referenzfakt. "
                    f"Bekannte Personen: {', '.join(known)}"
                )
            inj_l = injected_span.lower()
            if any(n.lower() in inj_l or inj_l in n.lower() for n in known):
                issues.append(
                    f"injected_span '{injected_span}' ist selbst ein Referenzfakt"
                )

    return issues


# ═══════════════════════════════════════════════════════════════════════════════
# HAUPTORCHESTRATOR: _process_segment
# ═══════════════════════════════════════════════════════════════════════════════

def _process_segment(
        seg_idx: int,
        seg_text: str,
        selected_class: str,
        llm: OllamaClient,
        entity_whitelist: List[str],
        reference_facts: Optional[Dict],
        case_id: str,
        proposal_counter: int,
        rng: random.Random,
) -> Optional[Dict]:
    """
    Führt Phase 1 (Extraktion) + Phase 3 (Injektion) für eine gegebene
    Fehlerklasse auf einem Segment aus.
    """

    if selected_class == "FORMALES_REDAKTION":
        # Phase 1: Precheck – welche Subtypen sind im Segment strukturell möglich?
        formales_candidates = extract_candidates_formales(seg_text, rng)
        if not formales_candidates:
            return None
        eligible_subtypes = [c["subtype"] for c in formales_candidates]
        # Dict {subtype → candidate_meta} für Prompt-Anreicherung (v.a. DOPPELKONSONANT)
        formales_meta = {c["subtype"]: c for c in formales_candidates}
        # Phase 3: LLM wählt aus validierten Subtypen – kein Halluzinationsrisiko
        proposal = inject_formales_llm(
            segment_text=seg_text, llm=llm,
            case_id=case_id, seg_idx=seg_idx,
            proposal_counter=proposal_counter, rng=rng,
            eligible_subtypes=eligible_subtypes,
            formales_candidates_meta=formales_meta,
        )

    elif selected_class == "RECHEN_ARITHMETIK":
        candidates = extract_candidates_rechen(seg_text)
        if not candidates:
            return None
        cand = rng.choice(candidates)
        proposal = inject_rechen(cand, seg_text, case_id, seg_idx, proposal_counter)

    elif selected_class == "RECHT_ABSICHERUNG":
        candidates = extract_candidates_absicherung(seg_text)
        if not candidates:
            return None
        cand = rng.choice(candidates)
        proposal = inject_absicherung_llm(
            cand, seg_text, llm, case_id, seg_idx, proposal_counter)

    elif selected_class == "RECHT_TRENNUNG":
        candidates = extract_candidates_trennung(seg_text)
        if not candidates:
            return None
        cand = rng.choice(candidates)
        proposal = inject_trennung_llm(
            cand, seg_text, llm, case_id, seg_idx, proposal_counter)

    elif selected_class == "HYPO_INKONSISTENZ":
        candidates = extract_candidates_hypo(seg_text)
        if not candidates:
            return None
        cand = rng.choice(candidates)
        proposal = inject_hypo_llm(
            cand, seg_text, llm, case_id, seg_idx, proposal_counter)

    elif selected_class == "STRUKT_BEFUND_BESCHREIBUNG":
        # Phase 1: konkrete Spans extrahieren
        strukt_candidates = extract_candidates_strukt(seg_text, reference_facts)
        available_subtypes = [st for st, lst in strukt_candidates.items() if lst]
        if not available_subtypes:
            return None
        # Subtyp zufällig wählen, dann einen konkreten Kandidaten-Span wählen
        subtype = rng.choice(available_subtypes)
        candidate = rng.choice(strukt_candidates[subtype])
        candidate_span = candidate["span_text"]
        # Phase 3: LLM muss nur noch ersetzen, nicht mehr suchen
        proposal = inject_strukt_llm(
            subtype=subtype,
            candidate_span=candidate_span,
            segment_text=seg_text,
            entity_whitelist=entity_whitelist,
            reference_facts=reference_facts,
            llm=llm,
            case_id=case_id,
            seg_idx=seg_idx,
            proposal_counter=proposal_counter,
        )
    else:
        return None

    if proposal is None:
        return None

    # Validierung
    issues = validate_proposal(proposal, seg_text, reference_facts)
    proposal["validation_ok"]     = len(issues) == 0
    proposal["validation_issues"] = issues
    return proposal




def check_plausibility(
        proposal: Dict[str, Any],
        llm: OllamaClient,
) -> Dict[str, Any]:
    """
    Post-hoc Plausibilitätsprüfung via LLM.
    Bewertet ob der Fehler in einem forensischen Gutachten realistisch wirkt.
    Gibt das Proposal mit zusätzlichen Feldern zurück:
      plausibility_ok:    bool
      plausibility_score: "hoch" | "mittel" | "niedrig"
      plausibility_reason: str
    Proposals mit plausibility_ok=False werden auf status="rejected" gesetzt.

    Vorab-Checks (ohne LLM):
    - Identische Strings: kein LLM-Call, sofort rejected
    - Leerer injected_span: sofort rejected
    """
    # ── Vorab-Checks: kein LLM nötig wenn Strings identisch oder leer ────────
    original  = (proposal.get("original_span") or "").strip()
    injected  = (proposal.get("injected_span") or "").strip()

    if not injected:
        return {
            **proposal,
            "plausibility_ok":     False,
            "plausibility_score":  "niedrig",
            "plausibility_reason": "injected_span ist leer – kein Fehler injiziert",
            "status":              "rejected",
        }

    if original == injected:
        return {
            **proposal,
            "plausibility_ok":     False,
            "plausibility_score":  "niedrig",
            "plausibility_reason": (
                "original_span und injected_span identisch – LLM hat nichts geändert, "
                "Kandidatensatz enthält vermutlich kein umformulierbares Sprachmuster"
            ),
            "status": "rejected",
        }

    system = (
        "Du bist ein Experte für forensische Gutachten und bewertest synthetische Fehler.\n"
        "Beurteile ob der beschriebene Fehler in einem echten forensischen Gutachten realistisch ist:\n"
        "- Könnte ein Mensch diesen Fehler tatsächlich machen?\n"
        "- Wirkt der injizierte Wert plausibel (nicht offensichtlich maschinell)?\n"
        "- Ist die Änderung minimal und nicht sofort als Artefakt erkennbar?\n"
        "Antworte NUR mit JSON:\n"
        "{\"plausibel\": bool, \"score\": \"hoch\"|\"mittel\"|\"niedrig\", \"grund\": string}"
    )
    user = (
        f"Fehlerklasse: {proposal.get('subclass_id')} / {proposal.get('subtype')}\n"
        f"Original: \"{proposal.get('original_span', '')}\"\n"
        f"Injiziert: \"{proposal.get('injected_span', '')}\"\n"
        f"Rationale: {proposal.get('rationale', '')}\n\n"
        "Ist dieser Fehler in einem forensischen Gutachten realistisch?"
    )

    schema = {
        "type": "object",
        "properties": {
            "plausibel": {"type": "boolean"},
            "score":     {"type": "string", "enum": ["hoch", "mittel", "niedrig"]},
            "grund":     {"type": "string"},
        },
        "required": ["plausibel", "score", "grund"],
    }

    try:
        raw = _llm_call(
            llm,
            [{"role": "system", "content": system}, {"role": "user", "content": user}],
            label="Plausibilitätsprüfung",
            schema=schema,)
        result = _parse_llm_json(raw)
        if not result:
            raise ValueError("Kein JSON")
    except Exception as e:
        # Bei Fehler: konservativ als plausibel markieren (kein false reject)
        return {
            **proposal,
            "plausibility_ok":     True,
            "plausibility_score":  "mittel",
            "plausibility_reason": f"Prüfung fehlgeschlagen: {e}",
        }

    plausibel = result.get("plausibel", True)
    score     = result.get("score", "mittel")
    grund     = result.get("grund", "").strip()

    updated = {
        **proposal,
        "plausibility_ok":     plausibel,
        "plausibility_score":  score,
        "plausibility_reason": grund,
    }
    # Unplausible Proposals automatisch auf rejected setzen
    if not plausibel:
        updated["status"] = "rejected"

    return updated

# ═══════════════════════════════════════════════════════════════════════════════
# MODUS 0: EXTRACT_FACTS
# ═══════════════════════════════════════════════════════════════════════════════

def run_extract_facts(args: argparse.Namespace) -> None:
    print("\n[INFO] Modus: EXTRACT_FACTS")
    doc_path = _load_docx_arg(args)
    doc_text = normalize_text(read_docx(doc_path))
    if not doc_text.strip():
        sys.exit("[FATAL] Kein Text extrahiert.")
    print(f"[INFO] Dokumenttext: {len(doc_text)} Zeichen")

    ref_schema_path = Path(env_str("REFERENCE_FACTS_SCHEMA",
                                   "schema/reference_facts_schema.json")).expanduser()
    if not ref_schema_path.exists():
        sys.exit(f"[FATAL] REFERENCE_FACTS_SCHEMA nicht gefunden: {ref_schema_path}")

    ref_schema = load_reference_facts_schema(ref_schema_path)
    llm        = make_injection_client()
    print("[INFO] Agent 0 (Referenzfakten-Extraktor) läuft...")

    try:
        reference_facts = run_reference_facts_agent(
            llm=llm, doc_text=doc_text, case_id=args.case_id,
            schema=ref_schema,
            max_chars=env_int("REFERENCE_FACTS_CONTEXT_CHARS", 4000),
        )
    except Exception as e:
        sys.exit(f"[FATAL] Extraktion fehlgeschlagen: {e}")

    out_path = _ref_facts_path(args.case_id, args.output_dir)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(reference_facts, ensure_ascii=False, indent=2),
                        encoding="utf-8")

    known = _known_person_names(reference_facts)
    print(f"[INFO] Personen (hoch/mittel): {', '.join(known) or '–'}")
    print(f"[INFO] Gespeichert: {out_path}")
    print(f"\nNÄCHSTER SCHRITT:")
    print(f"  python synthetic_error_injector.py --mode generate --document {args.document}")


# ═══════════════════════════════════════════════════════════════════════════════
# MODUS 1: GENERATE
# ═══════════════════════════════════════════════════════════════════════════════

def run_generate(args: argparse.Namespace) -> None:
    print("\n[INFO] Modus: GENERATE")

    doc_path = _load_docx_arg(args)
    doc_text = normalize_text(read_docx(doc_path))
    if not doc_text.strip():
        sys.exit("[FATAL] Kein Text extrahiert.")
    print(f"[INFO] Dokumenttext: {len(doc_text)} Zeichen")

    # Taxonomie-Validierung
    taxonomy_path = Path(env_str("TAXONOMY_JSON", "taxonomy.json")).expanduser().resolve()
    catalog = load_taxonomy_json(taxonomy_path)
    print(f"[INFO] Taxonomie: {len(catalog.main_classes)} Hauptklassen, "
          f"{sum(len(v) for v in catalog.sub_by_main.values())} Subklassen")

    segments = split_document_into_segments(
        doc_text,
        target_chars=env_int("SEG_TARGET_CHARS", 1200),
        min_chars=env_int("SEG_MIN_CHARS", 250),
        max_chars=env_int("SEG_MAX_CHARS", 2200),
    )
    print(f"[INFO] Segmente: {len(segments)}")

    # Entity-Whitelist
    whitelist_raw = env_str("INJECTION_ENTITY_WHITELIST", "")
    entity_whitelist: List[str]
    if whitelist_raw:
        try:
            entity_whitelist = json.loads(whitelist_raw)
        except json.JSONDecodeError:
            print("[WARN] INJECTION_ENTITY_WHITELIST ungültig – Fallback")
            entity_whitelist = _DEFAULT_ENTITY_WHITELIST
    else:
        entity_whitelist = _DEFAULT_ENTITY_WHITELIST

    target_errors = env_int("INJECTION_TARGET_ERRORS", 10)
    seed          = env_int("INJECTION_SEED", 42)
    case_id       = args.case_id
    rng           = random.Random(seed)

    print(f"[INFO] Case ID: {case_id} | Ziel-Fehler: {target_errors} | Seed: {seed}")
    print(f"[INFO] Fehlerklassen-Budget: "
          + ", ".join(f"{k}={int(v*100)}%" for k, v in CLASS_BUDGET.items()))

    # Referenzfakten laden (aus extract_facts, nie inline generiert)
    reference_facts: Optional[Dict[str, Any]] = None
    ref_facts_path = _ref_facts_path(case_id, args.output_dir)
    if ref_facts_path.exists():
        try:
            reference_facts = json.loads(ref_facts_path.read_text(encoding="utf-8"))
            known = _known_person_names(reference_facts)
            print(f"[INFO] Referenzfakten: {ref_facts_path.name}")
            print(f"[INFO] Personen ({len(known)}): {', '.join(known) or '–'}")
        except Exception as e:
            print(f"[WARN] Referenzfakten nicht lesbar: {e}")
    else:
        print(f"[WARN] Referenzfakten fehlen: {ref_facts_path}")
        print(f"[WARN] → python synthetic_error_injector.py "
              f"--mode extract_facts --document {args.document}")

    # Aktive Klassen (PERSON deaktiviert ohne Referenzfakten)
    allowed_classes = list(CLASS_PRIORITY)
    if reference_facts is None:
        allowed_classes = [c for c in allowed_classes if c != "STRUKT_BEFUND_BESCHREIBUNG"
                           or True]   # STRUKT bleibt, aber PERSON-Subtyp wird intern deaktiviert
        print("[INFO] PERSON-Subtyp deaktiviert (keine Referenzfakten)")

    # Pro Segment: prüfen welche Klassen möglich sind
    def available_classes_for_segment(seg: str) -> List[str]:
        classes = []
        if extract_candidates_formales(seg, rng):
            classes.append("FORMALES_REDAKTION")
        if extract_candidates_rechen(seg):
            classes.append("RECHEN_ARITHMETIK")
        if extract_candidates_absicherung(seg):
            classes.append("RECHT_ABSICHERUNG")
        if extract_candidates_trennung(seg):
            classes.append("RECHT_TRENNUNG")
        if extract_candidates_hypo(seg):
            classes.append("HYPO_INKONSISTENZ")
        strukt = extract_candidates_strukt(seg, reference_facts)
        if any(strukt.values()):
            classes.append("STRUKT_BEFUND_BESCHREIBUNG")
        return classes

    # Alle Segmente mit Kandidaten sammeln
    eligible: List[Tuple[int, str, List[str]]] = []
    for idx, seg in enumerate(segments):
        avail = available_classes_for_segment(seg)
        if avail:
            eligible.append((idx, seg, avail))

    print(f"[INFO] {len(eligible)}/{len(segments)} Segmente haben Kandidaten")
    if not eligible:
        sys.exit("[WARN] Keine geeigneten Segmente – Abbruch.")

    # Segmente samplen
    sampled = rng.sample(eligible, min(target_errors, len(eligible)))

    llm = make_injection_client()
    print(f"[INFO] Modell: {env_str('INJECTION_MODEL', '?')} | "
          f"Temp: {env_str('INJECTION_TEMPERATURE', '0.3')}")

    proposals:         List[Dict[str, Any]] = []
    produced_counts:   Dict[str, int]       = {}   # Klasse → Anzahl
    produced_subtypes: Dict[str, int]       = {}   # 'Klasse|Subtyp' → Anzahl
    seen_spans:        set                  = set()
    proposal_counter = 1

    # max_per_subtype: max. 1 Proposal pro Klasse|Subtyp-Kombination.
    # Verhindert monotone Wiederholungen (z.B. dreimal 'Datum in Zukunft').
    # Konfigurierbar via INJECTION_MAX_PER_SUBTYPE (default: 1).
    MAX_PER_SUBTYPE = env_int("INJECTION_MAX_PER_SUBTYPE", 1)

    for seg_idx, seg_text, avail_classes in sampled:
        print(f"\n[SEG {seg_idx:04d}] {seg_text[:80].replace(chr(10),' ')}...")
        print(f"  Verfügbar: {', '.join(avail_classes)}")

        selected = select_class(
            available_classes = avail_classes,
            produced_counts   = produced_counts,
            total_produced    = len(proposals),
            allowed_classes   = allowed_classes,
        )
        if selected is None:
            print("  [SKIP] Kein Budget für verfügbare Klassen")
            continue
        print(f"  Fehlerklasse: {selected}")

        proposal = _process_segment(
            seg_idx=seg_idx, seg_text=seg_text,
            selected_class=selected, llm=llm,
            entity_whitelist=entity_whitelist,
            reference_facts=reference_facts,
            case_id=case_id, proposal_counter=proposal_counter,
            rng=rng,
        )

        if proposal is None:
            print(f"  [SKIP] Kein Kandidat für {selected}")
            continue

        # ── Subtyp-Cap ────────────────────────────────────────────────────
        subtype     = proposal.get("subtype", "")
        subtype_key = f"{selected}|{subtype}"
        if produced_subtypes.get(subtype_key, 0) >= MAX_PER_SUBTYPE:
            print(
                f"  [SUBTYP-CAP] '{subtype_key}' bereits {MAX_PER_SUBTYPE}× "
                f"vorhanden – übersprungen"
            )
            continue

        # ── Duplikat-Filter (original_span) ───────────────────────────────
        orig = proposal.get("original_span", "").strip()
        if orig and orig in seen_spans:
            print(f"  [DUPLIKAT] '{orig[:60]}' – übersprungen")
            continue
        if orig:
            seen_spans.add(orig)

        sym = "✓" if proposal["validation_ok"] else "✗"
        print(f"  {sym} original : '{orig[:80]}'")
        print(f"  {sym} injiziert: '{proposal['injected_span'][:80]}'")
        if proposal.get("requires_human_review"):
            print(f"  [!] MANUELLE PRÜFUNG ERFORDERLICH")
        if proposal["validation_issues"]:
            for iss in proposal["validation_issues"]:
                print(f"  [VAL] ✗ {iss}")

        # Post-hoc Plausibilitätsprüfung
        proposal = check_plausibility(proposal, llm)
        plaus_sym = "✓" if proposal["plausibility_ok"] else "✗"
        print(f"  {plaus_sym} Plausibilität: {proposal['plausibility_score']} – {proposal['plausibility_reason'][:80]}")
        if not proposal["plausibility_ok"]:
            print(f"  [REJECT] Automatisch abgelehnt – nicht plausibel")

        proposals.append(proposal)
        produced_counts[selected]      = produced_counts.get(selected, 0) + 1
        produced_subtypes[subtype_key] = produced_subtypes.get(subtype_key, 0) + 1
        proposal_counter += 1
        _write_proposals(proposals, case_id, args.output_dir)

    # Zusammenfassung
    ok_count      = sum(1 for p in proposals if p.get("validation_ok"))
    plaus_ok      = sum(1 for p in proposals if p.get("plausibility_ok", True))
    auto_rejected = sum(1 for p in proposals if p.get("status") == "rejected")

    sep = "=" * 60
    summary_lines = [
        f"\n{sep}",
        f"Generiert:         {len(proposals)} Vorschläge",
        f"Valide (Validierung): {ok_count}/{len(proposals)}",
        f"Plausibel (LLM):     {plaus_ok}/{len(proposals)} ({auto_rejected} auto-rejected)",
        "Klassen-Verteilung (Ist / Ziel):",
    ]
    for cls in CLASS_PRIORITY:
        n   = produced_counts.get(cls, 0)
        tgt = int(CLASS_BUDGET.get(cls, 0) * 100)
        ist = int(n / max(len(proposals), 1) * 100)
        bar = "█" * n
        summary_lines.append(f"  {cls:<34} {n:2d}x  {ist:3d}% ist / {tgt:3d}% Ziel  {bar}")
    summary_lines.append("Subtyp-Verteilung:")
    for key in sorted(produced_subtypes):
        summary_lines.append(f"  {key:<48} {produced_subtypes[key]}x")

    path = _write_proposals(proposals, case_id, args.output_dir)
    summary_lines.append(f"Gespeichert: {path}")

    # Auf Konsole ausgeben
    for line in summary_lines:
        print(line)

    # In logs/ speichern
    log_dir  = (Path(__file__).parent / "logs").resolve()
    log_dir.mkdir(parents=True, exist_ok=True)
    ts_log   = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_path = log_dir / f"generate_{case_id}_{ts_log}.txt"
    with log_path.open("w", encoding="utf-8") as lf:
        lf.write(f"synthetic_error_injector.py --mode generate\n")
        lf.write(f"Zeitstempel: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        lf.write(f"Case: {case_id}\n")
        lf.write(f"Dokument: {args.document}\n")
        lf.write("\n")
        lf.write("\n".join(summary_lines).lstrip("\n"))
        lf.write("\n")
    print(f"Log gespeichert: {log_path}")

    print("\nNÄCHSTE SCHRITTE:")
    print(f"  1. {path.name} öffnen → status='accepted'/'rejected' setzen")
    print("   2. Fehler ins Dokument schreiben: --mode inject ausführen")
    print("   3. Ground Truth generieren: --mode validate ausführen")



# ═══════════════════════════════════════════════════════════════════════════════
# MODUS 2b: INJECT  —  akzeptierte Proposals ins Dokument schreiben
# ═══════════════════════════════════════════════════════════════════════════════
#
# Strategie: python-docx Run-aware Replacement
# Für jeden Paragraphen (inkl. Tabellenzellen) werden die Runs zu einem
# einzigen Text zusammengefügt, das Replacement durchgeführt, und der
# erste Run erhält den neuen Text. Alle anderen Runs des Paragraphen
# werden geleert. Run-interne Zeichenformatierung (z.B. einzelne fett-
# gedruckte Wörter innerhalb eines Runs) geht dabei verloren – die
# Absatzformatierung (Schrift, Grösse, Stil) bleibt erhalten.
#
# HINWEIS zu Reformulierungen (RECHT_ABSICHERUNG, RECHT_TRENNUNG, HYPO):
# Wenn original_span ein ganzer Satz ist, wird er ebenfalls über diesen
# Mechanismus ersetzt. Das funktioniert zuverlässig, solange der Satz
# in einem einzigen Paragraphen liegt (Normalfall). Sätze über mehrere
# Paragraphen werden nicht gefunden und im Log als SKIP markiert.

def _iter_all_paragraphs(doc: Any):
    """
    Generator: liefert alle Paragraphen eines python-docx Document-Objekts,
    inklusive Paragraphen in Tabellenzellen und verschachtelten Tabellen.
    Notwendig weil doc.paragraphs nur den Fliesstext sieht, aber Gutachten
    häufig Tabellen für Metadaten (Auftragsdaten, Personen) verwenden.
    """
    # Fliesstext-Paragraphen
    for para in doc.paragraphs:
        yield para
    # Tabellenzellen (alle Tiefen, rekursiv)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _norm(text: str) -> str:
    """
    Normalisiert einen String für den Span-Vergleich.
    Behandelt die häufigsten Quellen von Nicht-Übereinstimmungen zwischen
    dem normalisierten Segmentierungstext und dem XML-Rohtext eines .docx:

    - Non-breaking spaces (U+00A0) → normales Leerzeichen
      Word fügt diese automatisch ein (z.B. vor Einheiten, nach Abkürzungen).
    - Soft hyphens (U+00AD) → entfernen
      Silbentrennungsartefakte aus der Textverarbeitung.
    - Zeilenumbrüche (\n, \r) → normales Leerzeichen
      Paragraphen können interne Zeilenumbrüche enthalten (manueller
      Zeilenumbruch Shift+Enter in Word = <w:br/> im XML). Der
      Segmentierungstext behandelt diese als Leerzeichen; der XML-Rohtext
      enthält das Zeilenumbruch-Zeichen. Ohne diese Normalisierung schlägt
      der Vergleich fehl obwohl der Text identisch ist.
    - Unicode NFC-Normalisierung
      Stellt sicher dass zusammengesetzte Zeichen (ä = a + combining umlaut)
      und präkomponierte Zeichen (ä als einzelnes Zeichen) gleich behandelt werden.
    - Mehrfache Whitespace-Zeichen → einzelnes Leerzeichen
      Entsteht beim Zusammenführen von Runs, manueller Eingabe oder nach
      der Zeilenumbruch-Normalisierung.
    """
    text = text.replace('\u00a0', ' ')    # non-breaking space
    text = text.replace('\u00ad', '')     # soft hyphen
    text = text.replace('\r\n', ' ')      # Windows-Zeilenumbruch
    text = text.replace('\n', ' ')        # Unix-Zeilenumbruch / <w:br/>
    text = text.replace('\r', ' ')        # altes Mac-Zeilenumbruch
    text = unicodedata.normalize('NFC', text)
    text = re.sub(r'\s{2,}', ' ', text)   # mehrfache Whitespace → eines
    return text


def _replace_in_paragraph(para: Any, original: str, injected: str) -> bool:
    """
    Ersetzt original durch injected in einem einzelnen Paragraphen.
    Führt alle Runs zusammen, macht das Replacement, schreibt in Run 0,
    leert Runs 1..n.

    Gibt True zurück wenn eine Ersetzung stattfand, sonst False.

    Normalisierung: Der Vergleich erfolgt auf normalisierten Strings
    (_norm), das Replacement auf dem Rohtext. So werden Non-breaking
    Spaces, Soft-Hyphens und Unicode-Varianten korrekt gefunden, ohne
    den Original-XML zu verfälschen.

    Formatierungsverhalten:
    - Absatzformatierung (pPr) bleibt vollständig erhalten.
    - Zeichenformatierung des ersten Runs (rPr: Schrift, Grösse, Farbe)
      wird auf den gesamten neuen Text angewendet.
    - Run-interne Formatierungsunterschiede (z.B. ein einzelnes fettes Wort
      innerhalb des Paragraphen) gehen verloren, weil alle Runs zu einem
      zusammengeführt werden.
    """
    if not para.runs:
        return False

    full_text_raw  = "".join(r.text for r in para.runs)
    full_text_norm = _norm(full_text_raw)
    original_norm  = _norm(original)

    if original_norm not in full_text_norm:
        return False

    # Replacement auf Rohtext: Position via normalisiertem Index berechnen
    # Da _norm keine Zeichen hinzufügt (nur entfernt/ersetzt), ist eine
    # direkte Indexübertragung nicht trivial. Sicherste Strategie: Rohtext
    # ebenfalls normalisieren, dann ersetzen und original_span mit injected
    # ersetzen. Die Normalisierung des Rohtexts ist für die finale Ausgabe
    # akzeptabel, da sie nur unsichtbare/transparente Artefakte bereinigt.
    new_text = full_text_norm.replace(original_norm, injected, 1)
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""
    return True


def run_inject(args: argparse.Namespace) -> None:
    """
    Modus inject: Liest akzeptierte Proposals, kopiert das Originaldokument
    und schreibt die Fehler per Run-aware Replacement ins kopierte Dokument.

    Aufruf:
        python synthetic_error_injector.py --mode inject \
            --document ./case_documents/case_02.docx \
            --proposals injection/proposals_case_02_latest.json

    Output: case_documents/case_02_modified.docx

    Behandlung von Randfällen:
    - Wenn original_span im Dokument mehrfach vorkommt: nur das erste
      Vorkommen wird ersetzt (str.replace(..., 1)).
    - Wenn original_span gar nicht gefunden wird: SKIP-Meldung im Log,
      kein Fehler, Proposal bleibt unverändert.
    - Reformulierungen (is_reformulation=True): gleiche Logik, kein
      Sonderfall – der ganze Satz als original_span wird ersetzt.
    - Proposals mit status != "accepted": werden ignoriert.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        sys.exit(
            "[FATAL] python-docx nicht installiert.\n"
            "Bitte ausführen: pip install python-docx"
        )

    print("\n[INFO] Modus: INJECT")

    # ── Proposals laden ──────────────────────────────────────────────────────
    proposals_path = Path(args.proposals).expanduser().resolve()
    if not proposals_path.exists():
        sys.exit(f"[FATAL] Proposals-File nicht gefunden: {proposals_path}")

    with proposals_path.open(encoding="utf-8") as f:
        all_proposals: List[Dict[str, Any]] = json.load(f)

    accepted = [
        p for p in all_proposals
        if p.get("status") == "accepted"
           and (p.get("original_span") or "").strip()
           and (p.get("injected_span") or "").strip()
    ]
    print(f"[INFO] Proposals gesamt: {len(all_proposals)} | akzeptiert: {len(accepted)}")

    if not accepted:
        sys.exit("[WARN] Keine akzeptierten Proposals mit gültigen Spans – nichts zu tun.")

    # ── Originaldokument prüfen ──────────────────────────────────────────────
    doc_path = _load_docx_arg(args)

    # ── Ausgabepfad bestimmen: case_documents/case_02_modified.docx ─────────
    # Liegt immer im gleichen Verzeichnis wie das Originaldokument.
    # Suffix _modified wird vor der Erweiterung eingefügt.
    stem     = doc_path.stem                               # "case_02"
    out_name = f"{stem}_modified{doc_path.suffix}"         # "case_02_modified.docx"
    out_path = doc_path.parent / out_name                  # ".../case_documents/case_02_modified.docx"

    if out_path.exists():
        print(f"[WARN] Ausgabedatei existiert bereits und wird überschrieben: {out_path.name}")

    # ── Dokument kopieren ────────────────────────────────────────────────────
    # Wir arbeiten auf einer Kopie – das Original bleibt unverändert.
    shutil.copy2(doc_path, out_path)
    print(f"[INFO] Dokument kopiert: {doc_path.name} → {out_path.name}")

    # ── python-docx öffnen ───────────────────────────────────────────────────
    doc = DocxDocument(out_path)

    # ── Proposals anwenden ───────────────────────────────────────────────────
    n_ok   = 0
    n_skip = 0
    n_dup  = 0

    for proposal in accepted:
        original = proposal["original_span"].strip()
        injected = proposal["injected_span"].strip()
        pid      = proposal.get("proposal_id", "?")
        subtype  = proposal.get("subtype", "?")

        if original == injected:
            print(f"  [SKIP] {pid}: original_span == injected_span – übersprungen")
            n_skip += 1
            continue

        replaced = False
        for para in _iter_all_paragraphs(doc):
            if _replace_in_paragraph(para, original, injected):
                replaced = True
                break   # max. 1 Ersetzung pro Proposal (erstes Vorkommen)

        if replaced:
            print(f"  [OK]  {pid} ({subtype}): '{original[:50]}' → '{injected[:50]}'")
            n_ok += 1
        else:
            print(
                f"  [SKIP] {pid} ({subtype}): '{original[:50]}' "
                f"nicht im Dokument gefunden"
            )
            n_skip += 1

            # ── Diagnose-Logging bei SKIP (aktivieren via INJECT_DEBUG=true) ──
            if env_bool("INJECT_DEBUG", False):
                original_norm = _norm(original)
                partial       = original_norm[:40]
                matches       = []
                for i, para in enumerate(_iter_all_paragraphs(doc)):
                    para_norm = _norm("".join(r.text for r in para.runs))
                    if partial in para_norm:
                        matches.append((i, para_norm))

                if matches:
                    # Partieller Match → Paragraphengrenz-Problem
                    print(f"    [DIAGNOSE] Partieller Match '{partial[:30]}' in {len(matches)} Absatz/Absätzen:")
                    for idx, ptext in matches[:3]:
                        snip_start = max(0, ptext.index(partial) - 20)
                        snip_end   = ptext.index(partial) + 60
                        print(f"      Para #{idx}: ...{ptext[snip_start:snip_end]}...")
                    if len(matches) > 1:
                        print("    → Span vermutlich über Paragraphengrenze aufgeteilt")
                else:
                    # Kein partieller Match → Sonderzeichen / Unicode-Problem
                    print(f"    [DIAGNOSE] Kein partieller Match für '{partial[:30]}'")
                    print(f"    → original_span repr: {repr(original[:80])}")
                    # Erstes Wort suchen um Vorkommen im Dokument zu lokalisieren
                    words = original_norm.split()
                    first_word = words[0] if words else ""
                    if first_word:
                        for i, para in enumerate(_iter_all_paragraphs(doc)):
                            para_norm = _norm("".join(r.text for r in para.runs))
                            if first_word in para_norm:
                                w_idx = para_norm.index(first_word)
                                print(
                                    f"    → Erstes Wort '{first_word}' in Para #{i}: "
                                    f"...{para_norm[max(0,w_idx-10):w_idx+50]}..."
                                )
                                break

    # ── Speichern ────────────────────────────────────────────────────────────
    doc.save(out_path)

    print(f"\n{'='*60}")
    print(f"Injiziert: {n_ok} OK | {n_skip} nicht gefunden")
    print(f"Gespeichert: {out_path}")
    print()
    print("NÄCHSTER SCHRITT:")
    print(
        f"  python synthetic_error_injector.py --mode validate \\\n"
        f"      --proposals {proposals_path} \\\n"
        f"      --modified_doc {out_path}"
    )

# ═══════════════════════════════════════════════════════════════════════════════
# MODUS 2: VALIDATE
# ═══════════════════════════════════════════════════════════════════════════════

def run_validate(args: argparse.Namespace) -> None:
    print("\n[INFO] Modus: VALIDATE")

    proposals_path = Path(args.proposals).expanduser().resolve()
    if not proposals_path.exists():
        sys.exit(f"[FATAL] Proposals-File nicht gefunden: {proposals_path}")

    with proposals_path.open(encoding="utf-8") as f:
        proposals: List[Dict[str, Any]] = json.load(f)

    accepted = [p for p in proposals if p.get("status") == "accepted"]
    print(f"[INFO] Proposals: {len(proposals)} gesamt | {len(accepted)} akzeptiert")

    human_review = [p for p in accepted if p.get("requires_human_review")]
    if human_review:
        print(f"[WARN] {len(human_review)} Proposal(s) mit requires_human_review=true!")
        for p in human_review:
            print(f"  → {p['proposal_id']} ({p['subclass_id']})")

    if not accepted:
        sys.exit("[WARN] Keine akzeptierten Proposals.")

    mod_doc_path = Path(args.modified_doc).expanduser().resolve()
    if not mod_doc_path.exists():
        sys.exit(f"[FATAL] Modifiziertes Dokument nicht gefunden: {mod_doc_path}")

    print(f"[INFO] Modifiziertes Dokument: {mod_doc_path.name}")
    mod_text = normalize_text(read_docx(mod_doc_path))
    mod_segments = split_document_into_segments(
        mod_text,
        target_chars=env_int("SEG_TARGET_CHARS", 1200),
        min_chars=env_int("SEG_MIN_CHARS", 250),
        max_chars=env_int("SEG_MAX_CHARS", 2200),
    )
    print(f"[INFO] Segmente im modifizierten Dokument: {len(mod_segments)}")

    case_id    = args.case_id
    _p         = Path(args.output_dir).expanduser()
    output_dir = _p if _p.is_absolute() else (Path(__file__).parent / _p).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    # ── Segmentindex-Lookup-Tabelle: injected_span → tatsächlicher Segment-Index ─
    # Das modifizierte Dokument kann nach der Injektion leicht verschobene
    # Segmentgrenzen haben (z.B. durch Komposita-Trennungen oder umgeschriebene
    # Sätze). Statt den Segment-Index blind aus dem Proposal zu übernehmen,
    # wird für jeden accepted Proposal gesucht, in welchem Segment des
    # modifizierten Dokuments der injected_span tatsächlich vorkommt.
    # Das stellt sicher, dass segment_index in der GT mit dem modifizierten
    # Dokument übereinstimmt und das Evaluation-Framework Predictions und
    # GT-Findings korrekt abgleichen kann.

    gt_entries: List[Dict[str, Any]] = []
    finding_counter = 1
    all_ok = True

    for proposal in accepted:
        original_seg_idx = proposal["segment_index"]   # aus dem Originaldokument
        original_span    = (proposal.get("original_span") or "").strip()
        injected_span    = (proposal.get("injected_span") or "").strip()

        # ── Segment-Index neu bestimmen ───────────────────────────────────────
        # Suche das Segment im modifizierten Dokument, das injected_span enthält.
        actual_seg_idx = None
        for i, seg in enumerate(mod_segments):
            if injected_span in seg:
                actual_seg_idx = i
                break

        if actual_seg_idx is None:
            # injected_span nicht in Segmenten → Volltext-Check als Fallback
            if injected_span in mod_text:
                print(
                    f"  [WARN] Proposal {proposal.get('proposal_id','?')}: "
                    f"injected_span im Volltext, aber in keinem Segment – "
                    f"Segmentierungsparameter prüfen"
                )
                # Kein GT-Eintrag möglich ohne Segmentzuordnung
                all_ok = False
                continue
            else:
                print(
                    f"  [FAIL] Proposal {proposal.get('proposal_id','?')}: "
                    f"injected_span '{injected_span[:60]}' nicht im Dokument gefunden"
                )
                all_ok = False
                continue

        # Logging: verschobener vs. ursprünglicher Segment-Index
        if actual_seg_idx != original_seg_idx:
            print(
                f"  [SHIFT] Proposal {proposal.get('proposal_id','?')}: "
                f"Segment-Index verschoben {original_seg_idx} → {actual_seg_idx} "
                f"(GT verwendet korrekten Index {actual_seg_idx})"
            )
        else:
            mod_seg = mod_segments[actual_seg_idx]
            if original_span and original_span in mod_seg:
                print(
                    f"  [WARN] SEG {actual_seg_idx}: original_span noch vorhanden – "
                    f"Injektion korrekt?"
                )
            else:
                print(
                    f"  [OK]  SEG {actual_seg_idx}: '{original_span[:50]}' → "
                    f"'{injected_span[:50]}'"
                )

        mod_seg    = mod_segments[actual_seg_idx]
        finding_id = f"GT-{case_id}-{finding_counter:04d}"
        finding_counter += 1

        gt_finding = {
            "finding_id":    finding_id,
            "subclass_id":   proposal.get("subclass_id", ""),
            "change_type_id":proposal.get("change_type_id", ""),
            "severity_id":   proposal.get("severity_id", "MEDIUM"),
            "span_text":     injected_span,
            "correction":    original_span,
            "rationale":     proposal.get("rationale", ""),
        }

        # GT-Eintrag unter dem korrekten Segment-Index sammeln
        existing = next((e for e in gt_entries if e["segment_index"] == actual_seg_idx), None)
        if existing:
            existing["gold_findings"].append(gt_finding)
        else:
            gt_entries.append({
                "case_id":       case_id,
                "segment_id":    f"{case_id}_seg_{actual_seg_idx:04d}",
                "segment_index": actual_seg_idx,
                "gold_findings": [gt_finding],
            })

    ts      = datetime.now().strftime("%Y%m%d_%H%M%S")
    # Ground Truth ins /ground_truth/ Verzeichnis (relativ zum Script-Root)
    gt_dir  = (Path(__file__).parent / "ground_truth").resolve()
    gt_dir.mkdir(parents=True, exist_ok=True)
    gt_path = gt_dir / f"ground_truth_{case_id}_synthetic.jsonl"
    with gt_path.open("w", encoding="utf-8") as f:
        for entry in sorted(gt_entries, key=lambda e: e["segment_index"]):
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")

    total = sum(len(e["gold_findings"]) for e in gt_entries)
    print(f"\n{'='*60}")
    print(f"Ground Truth: {total} Findings in {len(gt_entries)} Segmenten")
    print(f"Gespeichert:  {gt_path}")
    if not all_ok:
        print("[WARN] Einige Validierungen fehlgeschlagen.")


# ═══════════════════════════════════════════════════════════════════════════════
# HILFSFUNKTIONEN
# ═══════════════════════════════════════════════════════════════════════════════

def _load_docx_arg(args: argparse.Namespace) -> Path:
    if not args.document:
        sys.exit("[FATAL] --document <pfad> fehlt")
    p = Path(args.document).expanduser().resolve()
    if not p.exists():
        sys.exit(f"[FATAL] Dokument nicht gefunden: {p}")
    if p.suffix.lower() != ".docx":
        sys.exit(f"[FATAL] Nur .docx unterstützt: {p.suffix}")
    print(f"[INFO] Dokument: {p.name}")
    return p


def _ref_facts_path(case_id: str, output_dir_str: str) -> Path:
    p = Path(output_dir_str).expanduser()
    d = p if p.is_absolute() else (Path(__file__).parent / p).resolve()
    return d / f"reference_facts_{case_id}_original.json"


def _write_proposals(proposals: List[Dict], case_id: str, output_dir_str: str) -> Path:
    p = Path(output_dir_str).expanduser()
    d = p if p.is_absolute() else (Path(__file__).parent / p).resolve()
    d.mkdir(parents=True, exist_ok=True)
    path = d / f"proposals_{case_id}_latest.json"
    with path.open("w", encoding="utf-8") as f:
        json.dump(proposals, f, ensure_ascii=False, indent=2)
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# ARGPARSE & MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser(
        description="Synthetische Fehlerinjektion für forensische Gutachten",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument("--mode", choices=["extract_facts", "generate", "inject", "validate"],
                    default="generate",
                    help="extract_facts | generate | inject | validate")
    ap.add_argument("--document", type=str,
                    default=env_str("INJECTION_DOCUMENT_PATH", ""),
                    help="Pfad zum Quell-.docx")
    ap.add_argument("--proposals", type=str, default="",
                    help="[validate] Pfad zu proposals_*.json")
    ap.add_argument("--modified_doc", type=str,
                    default=env_str("INJECTION_MODIFIED_DOC", ""),
                    help="[validate] Pfad zum modifizierten .docx")
    ap.add_argument("--output_dir", type=str,
                    default=env_str("INJECTION_OUTPUT_DIR", "injection"),
                    help="Ausgabeverzeichnis (default: injection)")

    args = ap.parse_args()

    # case_id Prioritätsreihenfolge:
    # 1. CASE_ID aus .env (explizit gesetzt)
    # 2. Stem von --document  (generate / extract_facts / inject)
    # 3. case_id-Feld aus erstem Eintrag des --proposals File  (validate / inject)
    # 4. Fallback "case_unknown"
    env_case_id = env_str("CASE_ID", "")
    if env_case_id:
        args.case_id = env_case_id
    elif args.document:
        args.case_id = Path(args.document).stem
    elif args.proposals:
        try:
            p = Path(args.proposals).expanduser().resolve()
            if p.exists():
                import json as _j
                data = _j.loads(p.read_text(encoding="utf-8"))
                args.case_id = (data[0].get("case_id") or "case_unknown") if data else "case_unknown"
            else:
                args.case_id = "case_unknown"
        except Exception:
            args.case_id = "case_unknown"
    else:
        args.case_id = "case_unknown"

    return args


def main() -> None:
    t0 = time.perf_counter()
    print(f"[INFO] synthetic_error_injector.py – "
          f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    args = parse_args()

    if args.mode == "extract_facts":
        run_extract_facts(args)
    elif args.mode == "generate":
        run_generate(args)
    elif args.mode == "inject":
        if not args.proposals:
            sys.exit("[FATAL] --mode inject erfordert --proposals")
        if not args.document:
            sys.exit("[FATAL] --mode inject erfordert --document")
        run_inject(args)
    elif args.mode == "validate":
        if not args.proposals:
            sys.exit("[FATAL] --mode validate erfordert --proposals")
        if not args.modified_doc:
            sys.exit("[FATAL] --mode validate erfordert --modified_doc")
        run_validate(args)

    print(f"[INFO] Fertig in {time.perf_counter() - t0:.1f}s")


if __name__ == "__main__":
    main()