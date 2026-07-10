#!/usr/bin/env python3
"""
llm_error_detector.py
=====================
Erkennt echte Fehler in einem forensischen Gutachten mittels der bestehenden
RAG-Agenten-Pipeline (rag_answer_reference_facts.py) und stellt sie als
prüfbares JSON-File bereit, das vom Enduser akzeptiert oder abgelehnt werden kann.

Im Gegensatz zum synthetic_error_injector.py werden hier keine Fehler
synthetisch eingebaut, sondern echte Fehler im Originaldokument gesucht.
Der Workflow ist analog aufgebaut: detect → (human review) → apply / to_gt.

WORKFLOW (3 Schritte):
  1. python llm_error_detector.py --mode detect \\
         --document ./case_documents/case_xy.docx \\
         --case_id xy
     → detection/findings_case_xy_<ts>.json   (alle status="pending")
     → Enduser öffnet JSON, setzt status pro Finding:
         "confirmed"  – Fehler bestätigt, corrected_span korrekt
         "rejected"   – kein Fehler (LLM-Halluzination)
         "uncertain"  – unklar, noch nicht entschieden
         "corrected"  – Fehler bestätigt, aber corrected_span falsch
                        → Enduser füllt human_correction aus

  2. python llm_error_detector.py --mode apply \\
         --document ./case_documents/case_xy.docx \\
         --findings detection/findings_case_xy_<ts>.json
     → case_documents/case_xy_corrected.docx
     (Kopie des Originals mit angewandten Korrekturen für "confirmed" + "corrected")

  3. python llm_error_detector.py --mode to_gt \\
         --document ./case_documents/case_xy.docx \\
         --findings detection/findings_case_xy_<ts>.json
     → ground_truth/ground_truth_case_xy_real.jsonl
     (kompatibel mit evaluate_predictions.py)

FINDINGS-SCHEMA (detection/findings_case_xy_<ts>.json):
  {
    "finding_id":       "DET-case_xy-0001",
    "case_id":          "case_xy",
    "source":           "llm_detection",
    "status":           "pending",       # pending | confirmed | rejected | uncertain | corrected
    "agent":            "factual",       # welcher Agent hat das Finding erzeugt
    "segment_index":    3,
    "subclass_id":      "STRUKT_BEFUND_BESCHREIBUNG",
    "change_type_id":   "CHANGE_FACHLICH",
    "severity_id":      "MEDIUM",
    "original_span":    "15,3 km/h",    # was im Dokument steht (der potenzielle Fehler)
    "corrected_span":   "13,5 km/h",    # was da stehen sollte (LLM-Vorschlag)
    "rationale":        "...",
    "human_correction": ""              # Enduser füllt aus wenn corrected_span falsch ist
  }

ENV-VARIABLEN (.env):
  OLLAMA_BASE_URL, LLM_BACKEND, LLM_TIMEOUT_S,
  CASE_ID, TAXONOMY_JSON,
  REFERENCE_FACTS_SCHEMA, REFERENCE_FACTS_CONTEXT_CHARS,
  EMBED_MODEL, EMBED_DEVICE,
  DETECTION_OUTPUT_DIR   (default: detection)
  RAG_EMBEDDINGS, RAG_INDEX, RAG_PREPARED
  RAG_EMBEDDINGS2, RAG_INDEX2, RAG_PREPARED2
  RULES_TOP_K, MATERIAL_TOP_K, CONTEXT_MAX_CHARS
  MULTI_QUERY_COUNT, MMR_LAMBDA, MAX_PER_SOURCE
  VISION_MODEL, VISION_WORKERS, VISION_TIMEOUT_S
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import time
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


# ── Bestehende Infrastruktur importieren ──────────────────────────────────────
_THIS_DIR = Path(__file__).resolve().parent
if str(_THIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THIS_DIR))

try:
    import rag_answer_reference_facts as rag
    from rag_answer_reference_facts import (
        load_dotenv, env_str, env_int, env_bool,
        require_env, env_json_object_optional,
        load_taxonomy_json, ErrorCatalog,
        build_label_to_id_maps,
        LLMClient, make_llm_client,
        load_hf_model, choose_device,
        load_rag_store,
        check_document,
        load_reference_facts_schema,
        split_document_into_segments,
        _find_segment_for_stelle,
    )
except ImportError as e:
    sys.exit(
        f"[FATAL] rag_answer_reference_facts.py konnte nicht importiert werden:\n  {e}\n"
        "Sicherstellen dass das File im selben Verzeichnis liegt."
    )

try:
    from importDocuments_structural import normalize_text, read_docx
except ImportError as e:
    sys.exit(f"[FATAL] importDocuments_structural.py nicht gefunden:\n  {e}")

load_dotenv(".env")


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX-HILFSFUNKTIONEN  (analog zu synthetic_error_injector.py)
# ═══════════════════════════════════════════════════════════════════════════════

def _iter_all_paragraphs(doc: Any):
    """
    Generator über alle Paragraphen eines python-docx Dokuments,
    inklusive Paragraphen in Tabellenzellen (alle Tiefen).
    Notwendig weil doc.paragraphs nur den Fliesstext sieht, aber Gutachten
    häufig Tabellen für Metadaten (Auftragsdaten, Personen) verwenden.
    """
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _norm(text: str) -> str:
    """
    Normalisiert einen String für den Span-Vergleich im docx-XML.
    Behandelt Non-breaking Spaces, Soft Hyphens, Zeilenumbrüche,
    Unicode-Varianten und mehrfache Whitespace-Zeichen.
    """
    text = text.replace('\u00a0', ' ')   # non-breaking space
    text = text.replace('\u00ad', '')    # soft hyphen
    text = text.replace('\r\n', ' ')     # Windows-Zeilenumbruch
    text = text.replace('\n', ' ')       # Unix-Zeilenumbruch / <w:br/>
    text = text.replace('\r', ' ')       # altes Mac-Zeilenumbruch
    text = unicodedata.normalize('NFC', text)
    text = re.sub(r'\s{2,}', ' ', text)  # mehrfache Whitespace → eines
    return text


def _replace_in_paragraph(para: Any, original: str, replacement: str) -> bool:
    """
    Ersetzt original durch replacement in einem einzelnen Paragraphen.
    Führt alle Runs zusammen, macht das Replacement auf normalisierten Strings,
    schreibt in Run 0, leert Runs 1..n.

    Gibt True zurück wenn eine Ersetzung stattfand, sonst False.
    Nur das erste Vorkommen wird ersetzt (str.replace(..., 1)).
    """
    if not para.runs:
        return False

    full_text_norm = _norm("".join(r.text for r in para.runs))
    original_norm  = _norm(original)

    if original_norm not in full_text_norm:
        return False

    new_text = full_text_norm.replace(original_norm, replacement, 1)
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""
    return True


def _apply_replacements_to_docx(
        doc_path: Path,
        out_path: Path,
        replacements: List[Tuple[str, str, str]],
) -> Tuple[int, int]:
    """
    Kopiert doc_path nach out_path und wendet eine Liste von
    (finding_id, original_span, corrected_span)-Tupeln an.

    Gibt (n_ok, n_skip) zurück.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        sys.exit(
            "[FATAL] python-docx nicht installiert.\n"
            "Bitte ausführen: pip install python-docx"
        )

    shutil.copy2(doc_path, out_path)
    doc = DocxDocument(out_path)

    n_ok   = 0
    n_skip = 0

    for finding_id, original, corrected in replacements:
        if not original.strip() or not corrected.strip():
            print(f"  [SKIP] {finding_id}: leerer original_span oder corrected_span")
            n_skip += 1
            continue

        if _norm(original) == _norm(corrected):
            print(f"  [SKIP] {finding_id}: original_span == corrected_span – keine Änderung")
            n_skip += 1
            continue

        replaced = False
        for para in _iter_all_paragraphs(doc):
            if _replace_in_paragraph(para, original, corrected):
                replaced = True
                break  # max. 1 Ersetzung pro Finding (erstes Vorkommen)

        if replaced:
            print(f"  [OK]  {finding_id}: '{original[:60]}' → '{corrected[:60]}'")
            n_ok += 1
        else:
            print(f"  [SKIP] {finding_id}: '{original[:60]}' nicht im Dokument gefunden")
            n_skip += 1

    doc.save(out_path)
    return n_ok, n_skip


# ═══════════════════════════════════════════════════════════════════════════════
# FINDINGS-KONVERTIERUNG  (Report-Dict → proposals-kompatibles JSON)
# ═══════════════════════════════════════════════════════════════════════════════

# Mapping von internem Agent-Key auf menschenlesbaren Namen
_AGENT_LABELS: Dict[str, str] = {
    "factual":               "factual",
    "language":              "language",
    "calculation":           "calculation",
    "hypothesis":            "hypothesis",
    "reference_consistency": "reference_consistency",
    "statement_assurance":   "statement_assurance",
}

# Welches Feld enthält den Korrekturvorschlag pro Agent-Typ?
# Language-Agent hat "vorschlag", Calculation hat "korrekter_wert",
# alle anderen haben kein dediziertes Korrekturfeld.
_CORRECTION_FIELD: Dict[str, str] = {
    "language":    "vorschlag",
    "calculation": "korrekter_wert",
}


def _extract_corrected_span(item: Dict[str, Any], agent: str) -> str:
    """
    Extrahiert den LLM-Korrekturvorschlag aus einem Finding-Dict.

    Priorität:
    1. agent-spezifisches Feld ("vorschlag" für language, "korrekter_wert" für calculation)
    2. generisches "correction"-Feld (falls vorhanden)
    3. leerer String (Enduser muss manuell ausfüllen)
    """
    field = _CORRECTION_FIELD.get(agent)
    if field:
        val = str(item.get(field) or "").strip()
        if val:
            return val

    val = str(item.get("correction") or "").strip()
    if val:
        return val

    return ""


def convert_report_to_findings(
        report: Dict[str, Any],
        case_id: str,
        catalog: ErrorCatalog,
        doc_text: str,
        segments: List[str],
) -> List[Dict[str, Any]]:
    """
    Konvertiert das check_document()-Report-Dict in eine Liste von
    Finding-Dicts im proposals-kompatiblen Schema mit source="llm_detection".

    Jedes Finding erhält:
    - finding_id:       eindeutige ID (DET-<case_id>-NNNN)
    - case_id:          aus Argument
    - source:           "llm_detection"
    - status:           "pending"
    - agent:            welcher Agent das Finding erzeugt hat
    - segment_index:    aus dem Finding
    - subclass_id:      aus Taxonomie-Mapping (Label → ID)
    - change_type_id:   aus Taxonomie-Mapping
    - severity_id:      aus Taxonomie-Mapping
    - original_span:    stelle_im_segment (was im Dokument steht)
    - corrected_span:   LLM-Korrekturvorschlag (leer wenn nicht verfügbar)
    - rationale:        begruendung
    - human_correction: leer (Enduser füllt aus wenn corrected_span falsch)
    """
    sub_map, change_map, severity_map = build_label_to_id_maps(catalog)

    # Alle Finding-Listen aus dem Report mit Agent-Label zusammenführen
    agent_groups: List[Tuple[str, List[Dict[str, Any]]]] = [
        ("factual",               report.get("factual_findings", []) or []),
        ("language",              report.get("language_findings", []) or []),
        ("calculation",           report.get("calculation_findings", []) or []),
        ("hypothesis",            report.get("hypothesis_findings", []) or []),
        ("reference_consistency", report.get("reference_consistency_findings", []) or []),
        ("statement_assurance",   report.get("statement_assurance_findings", []) or []),
    ]

    findings: List[Dict[str, Any]] = []
    counter = 1

    for agent, items in agent_groups:
        for item in items:
            if not isinstance(item, dict):
                continue

            seg_idx = item.get("segment_index")
            if seg_idx is None:
                continue

            # original_span: Priorität stelle_im_segment, dann span_text
            original_span = str(
                item.get("stelle_im_segment") or item.get("span_text") or ""
            ).strip()
            if not original_span:
                print(
                    f"[CONVERT DROP] agent={agent} seg={seg_idx}: "
                    f"kein stelle_im_segment – übersprungen"
                )
                continue

            # Taxonomie-Labels → IDs
            subclass_label  = str(item.get("subklasse") or "").strip()
            change_label    = str(item.get("aenderungstyp") or "").strip()
            severity_label  = str(item.get("schweregrad") or "").strip()

            subclass_id    = sub_map.get(subclass_label)
            change_type_id = change_map.get(change_label)
            severity_id    = severity_map.get(severity_label)

            # Case-insensitive Fallback
            if not subclass_id:
                subclass_id = next(
                    (v for k, v in sub_map.items()
                     if k.strip().lower() == subclass_label.lower()),
                    None,
                )
            if not change_type_id:
                change_type_id = next(
                    (v for k, v in change_map.items()
                     if k.strip().lower() == change_label.lower()),
                    None,
                )
            if not severity_id:
                severity_id = next(
                    (v for k, v in severity_map.items()
                     if k.strip().lower() == severity_label.lower()),
                    None,
                )

            if not subclass_id or not change_type_id or not severity_id:
                print(
                    f"[CONVERT DROP] agent={agent} seg={seg_idx} "
                    f"subclass={subclass_label!r}→{subclass_id!r} "
                    f"change={change_label!r}→{change_type_id!r} "
                    f"severity={severity_label!r}→{severity_id!r} – übersprungen"
                )
                continue

            corrected_span = _extract_corrected_span(item, agent)

            finding: Dict[str, Any] = {
                "finding_id":       f"DET-{case_id}-{counter:04d}",
                "case_id":          case_id,
                "source":           "llm_detection",
                "status":           "pending",
                "agent":            agent,
                "segment_index":    seg_idx,
                "subclass_id":      subclass_id,
                "change_type_id":   change_type_id,
                "severity_id":      severity_id,
                "original_span":    original_span,
                "corrected_span":   corrected_span,
                "rationale":        str(item.get("begruendung") or "").strip(),
                "human_correction": "",
            }

            findings.append(finding)
            counter += 1

    return findings


# ═══════════════════════════════════════════════════════════════════════════════
# MODUS 1: DETECT
# ═══════════════════════════════════════════════════════════════════════════════


def _resolve_project_path(value: str) -> Path:
    """
    Löst Pfade robust relativ zum Projektordner auf.

    Hintergrund: Die GUI startet llm_error_detector.py zwar mit cwd=PROJECT_DIR,
    aber ein Wrapper sollte auch dann funktionieren, wenn er aus einem anderen
    Arbeitsverzeichnis gestartet wird.
    """
    p = Path(value).expanduser()
    if p.is_absolute():
        return p.resolve()
    return (_THIS_DIR / p).resolve()


def _ensure_parent(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def _default_detection_output_paths(args: argparse.Namespace) -> Dict[str, Path]:
    """
    Standardisierte Artefakte pro Case.

    Für case_14 entstehen damit u.a.:
      predictions/predictions_case_14.jsonl
      reference_facts/reference_facts_case_14.json
      reference_facts/reference_facts_case_14.jsonl
      reports/report_case_14.json
      detection/findings_case_14_<ts>.json
    """
    case_id = (args.case_id or "case_unknown").strip()

    predictions_path = (
        _resolve_project_path(args.save_predictions_jsonl)
        if getattr(args, "save_predictions_jsonl", "")
        else (_THIS_DIR / "predictions" / f"predictions_{case_id}.jsonl").resolve()
    )
    reference_json_path = (
        _resolve_project_path(args.save_reference_facts_json)
        if getattr(args, "save_reference_facts_json", "")
        else (_THIS_DIR / "reference_facts" / f"reference_facts_{case_id}.json").resolve()
    )
    reference_jsonl_path = reference_json_path.with_suffix(".jsonl")
    report_json_path = (
        _resolve_project_path(args.save_report_json)
        if getattr(args, "save_report_json", "")
        else (_THIS_DIR / "reports" / f"report_{case_id}.json").resolve()
    )

    return {
        "predictions": _ensure_parent(predictions_path),
        "reference_json": _ensure_parent(reference_json_path),
        "reference_jsonl": _ensure_parent(reference_jsonl_path),
        "report_json": _ensure_parent(report_json_path),
    }


def _write_reference_facts_jsonl_alias(reference_json_path: Path, reference_jsonl_path: Path) -> None:
    """
    rag_answer_reference_facts.py schreibt Referenzfakten als JSON-Objekt.
    Für den GUI-/Review-Workflow wird zusätzlich eine JSONL-Variante geschrieben,
    damit pro Case ein konsistentes Artefakt mit .jsonl-Endung vorhanden ist.
    """
    if not reference_json_path.exists():
        print(f"[WARN] Reference facts JSON nicht gefunden, JSONL-Alias übersprungen: {reference_json_path}")
        return
    try:
        obj = json.loads(reference_json_path.read_text(encoding="utf-8"))
    except Exception as exc:
        print(f"[WARN] Reference facts JSON konnte nicht gelesen werden: {exc}")
        return
    reference_jsonl_path.parent.mkdir(parents=True, exist_ok=True)
    reference_jsonl_path.write_text(
        json.dumps(obj, ensure_ascii=False) + "\n",
        encoding="utf-8",
        )
    print(f"[INFO] Saved reference facts JSONL to {reference_jsonl_path}")


def _run_rag_answer_reference_facts_pipeline(
        args: argparse.Namespace,
        *,
        artifact_paths: Dict[str, Path],
) -> Tuple[Dict[str, Any], ErrorCatalog, str, List[str]]:
    """
    Echter Wrapper um rag_answer_reference_facts.py.

    Wichtig: Hier wird keine Agentenlogik nachgebaut. Der Detector benutzt die
    zentrale check_document()-Pipeline aus rag_answer_reference_facts.py. Damit
    bleiben CLI und GUI fachlich auf demselben Codepfad.

    Das vollständige Report-Dict wird über einen schmalen Monkeypatch von
    save_predictions_jsonl() abgefangen. Die Originalfunktion wird trotzdem
    ausgeführt, sodass predictions_<case_id>.jsonl exakt wie im Hauptskript
    geschrieben wird.
    """
    doc_path = _resolve_project_path(args.document)
    if not doc_path.exists():
        sys.exit(f"[FATAL] Dokument nicht gefunden: {doc_path}")
    if doc_path.suffix.lower() != ".docx":
        sys.exit(f"[FATAL] Nur .docx unterstützt: {doc_path.suffix}")

    print(f"[INFO] Dokument: {doc_path.name}")
    print("[INFO] Starte Hauptpipeline: rag_answer_reference_facts.check_document()")

    # Dieselben Artefakte schreiben wie beim direkten Konsolenaufruf.
    args.save_predictions_jsonl = str(artifact_paths["predictions"])
    args.save_reference_facts_json = str(artifact_paths["reference_json"])

    # In der GUI sollen die gleichen Debug-Informationen im Log sichtbar sein
    # wie beim bisherigen Konsolenaufruf. Über ENV kann man das bei Bedarf dämpfen.
    if env_bool("DETECTION_PRINT_SOURCES", True):
        args.print_sources = True
    if env_bool("DETECTION_PRINT_CONTEXT", True):
        args.print_context = True
    if env_bool("DETECTION_PRINT_REFERENCE_FACTS", True):
        args.print_reference_facts = True

    catalog = rag.load_taxonomy_json(_resolve_project_path(args.taxonomy_json))
    print(
        f"[INFO] Taxonomie: {len(catalog.main_classes)} Hauptklassen | "
        f"{sum(len(v) for v in catalog.sub_by_main.values())} Subklassen | "
        f"{len(catalog.change_types)} Änderungstypen"
    )

    stores = [
        rag.load_rag_store(
            "rules",
            "rules",
            npz_path=_resolve_project_path(args.embeddings),
            index_path=_resolve_project_path(args.index),
            prepared_path=_resolve_project_path(args.prepared),
        )
    ]

    if str(args.embeddings2 or "").strip():
        if not str(args.index2 or "").strip() or not str(args.prepared2 or "").strip():
            sys.exit("[FATAL] RAG_EMBEDDINGS2 ist gesetzt, aber RAG_INDEX2 oder RAG_PREPARED2 fehlt")
        stores.append(
            rag.load_rag_store(
                "material",
                "material",
                npz_path=_resolve_project_path(args.embeddings2),
                index_path=_resolve_project_path(args.index2),
                prepared_path=_resolve_project_path(args.prepared2),
            )
        )
    else:
        print("[INFO] Kein zweiter RAG-Store konfiguriert — nur Normbasis.")

    device = rag.choose_device(args.embed_device)
    print(f"[INFO] Embedding device: {device} | model: {args.embed_model}")
    embed_model, embed_tok = rag.load_hf_model(args.embed_model, device)
    llm = rag.make_llm_client()

    vision_cfg: Optional[dict] = None
    vision_model_enabled = rag.env_bool("VISION_MODEL_ENABLED", True)
    if args.vision_model.strip() and vision_model_enabled:
        vision_cfg = {
            "vision_model": args.vision_model,
            "vision_model_enabled": vision_model_enabled,
            "vision_prompt": rag.env_str("VISION_PROMPT", ""),
            "ollama_base_url": rag.require_env("OLLAMA_BASE_URL"),
            "vision_timeout_s": args.vision_timeout_s,
            "vision_options": rag.env_json_object_optional("VISION_OPTIONS_JSON"),
        }
        print(f"[INFO] Vision captioning aktiv: {args.vision_model}")
    else:
        print(
            f"[INFO] Vision captioning deaktiviert: "
            f"vision_model={args.vision_model!r}, vision_model_enabled={vision_model_enabled}"
        )

    print(
        f"[INFO] Retrieval: multi_query_count={args.multi_query_count} | "
        f"mmr_lambda={args.mmr_lambda:.2f} | max_per_source={args.max_per_source} | "
        f"rules_top_k={args.rules_top_k} | material_top_k={args.material_top_k}"
    )

    captured: Dict[str, Any] = {}
    original_save_predictions = rag.save_predictions_jsonl

    def _capturing_save_predictions_jsonl(*, report: Dict[str, Any], case_id: str, output_path: Path, catalog: ErrorCatalog) -> None:
        captured["report"] = report
        original_save_predictions(
            report=report,
            case_id=case_id,
            output_path=output_path,
            catalog=catalog,
        )

    rag.save_predictions_jsonl = _capturing_save_predictions_jsonl  # type: ignore[assignment]
    try:
        rag.check_document(
            doc_path,
            stores,
            embed_model=embed_model,
            embed_tok=embed_tok,
            device=device,
            llm=llm,
            args=args,
            vision_cfg=vision_cfg,
            catalog=catalog,
        )
    finally:
        rag.save_predictions_jsonl = original_save_predictions  # type: ignore[assignment]

    report = captured.get("report")
    if not isinstance(report, dict):
        sys.exit(
            "[FATAL] rag_answer_reference_facts.py hat kein Report-Dict geliefert. "
            "Prüfe, ob save_predictions_jsonl() in check_document() noch aufgerufen wird."
        )

    report_json_path = artifact_paths["report_json"]
    report_json_path.write_text(
        json.dumps(report, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"[INFO] Saved full report JSON to {report_json_path}")

    _write_reference_facts_jsonl_alias(
        artifact_paths["reference_json"],
        artifact_paths["reference_jsonl"],
    )

    # Für die Segment-Verifikation beim GT-Export und für die Konvertierung.
    doc_text = normalize_text(read_docx(doc_path))
    segments = rag.split_document_into_segments(
        doc_text,
        target_chars=env_int("SEG_TARGET_CHARS", 1200),
        min_chars=env_int("SEG_MIN_CHARS", 250),
        max_chars=env_int("SEG_MAX_CHARS", 2200),
    )

    return report, catalog, doc_text, segments


def run_detect(args: argparse.Namespace) -> None:
    """
    DETECT ist jetzt ein echter Wrapper um rag_answer_reference_facts.py.

    Die fachliche Pipeline läuft ausschliesslich über check_document() aus
    rag_answer_reference_facts.py. llm_error_detector.py erzeugt danach nur
    noch das GUI-kompatible Review-JSON und bleibt für apply/to_gt zuständig.
    """
    print("\n[INFO] Modus: DETECT")

    artifact_paths = _default_detection_output_paths(args)
    report, catalog, doc_text, segments = _run_rag_answer_reference_facts_pipeline(
        args,
        artifact_paths=artifact_paths,
    )

    print("\n[INFO] Konvertiere Report ins Detection-/GUI-Schema ...")
    findings = convert_report_to_findings(
        report, args.case_id, catalog, doc_text, segments
    )

    total_raw = sum(
        len(report.get(k, []) or [])
        for k in [
            "factual_findings",
            "language_findings",
            "calculation_findings",
            "hypothesis_findings",
            "reference_consistency_findings",
            "statement_assurance_findings",
        ]
    )
    print(
        f"[INFO] Findings: {total_raw} roh aus rag_answer_reference_facts.py | "
        f"{len(findings)} nach Konvertierung ins GUI-Schema"
    )

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_dir = _resolve_output_dir(args.output_dir)
    out_path = out_dir / f"findings_{args.case_id}_{ts}.json"

    with out_path.open("w", encoding="utf-8") as f:
        json.dump(findings, f, ensure_ascii=False, indent=2)

    sep = "=" * 60
    print(f"\n{sep}")
    print(f"Findings gesamt: {len(findings)}")
    _print_agent_summary(findings)
    print(f"Predictions JSONL:        {artifact_paths['predictions']}")
    print(f"Reference facts JSON:     {artifact_paths['reference_json']}")
    print(f"Reference facts JSONL:    {artifact_paths['reference_jsonl']}")
    print(f"Full report JSON:         {artifact_paths['report_json']}")
    print(f"Gespeichert:     {out_path}")
    print(f"\nNÄCHSTE SCHRITTE:")
    print(f"  1. {out_path.name} in der GUI prüfen")
    print(f"     → status='confirmed'/'rejected'/'uncertain'/'corrected' setzen")
    print(f"     → bei status='corrected': human_correction ausfüllen")
    print(f"  2. Korrekturen anwenden:")
    print(
        f"     python llm_error_detector.py --mode apply \\\n"
        f"         --document {_resolve_project_path(args.document)} \\\n"
        f"         --findings {out_path}"
    )
    print(f"  3. Ground Truth erzeugen:")
    print(
        f"     python llm_error_detector.py --mode to_gt \\\n"
        f"         --document {_resolve_project_path(args.document)} \\\n"
        f"         --findings {out_path}"
    )


# ═══════════════════════════════════════════════════════════════════════════════
# MODUS 2: APPLY
# ═══════════════════════════════════════════════════════════════════════════════

def run_apply(args: argparse.Namespace) -> None:
    """
    Liest die vom Enduser geprüften Findings und wendet alle Korrekturen
    (status="confirmed" oder status="corrected") auf das Originaldokument an.

    - "confirmed": corrected_span wird als Korrektur verwendet
    - "corrected":  human_correction wird als Korrektur verwendet
                    (Enduser hat den LLM-Vorschlag überschrieben)

    Das Originaldokument bleibt unverändert; es wird eine neue Datei
    case_xy_corrected.docx im selben Verzeichnis erzeugt.
    """
    print("\n[INFO] Modus: APPLY")

    doc_path      = _load_doc_arg(args)
    findings_path = _load_findings_arg(args)
    findings      = _read_findings(findings_path)

    # Nur bestätigte Findings anwenden
    to_apply = [
        f for f in findings
        if f.get("status") in ("confirmed", "corrected")
    ]

    if not to_apply:
        print("[WARN] Keine bestätigten Findings (confirmed/corrected) – nichts zu tun.")
        return

    print(
        f"[INFO] Findings gesamt: {len(findings)} | "
        f"anzuwenden: {len(to_apply)} "
        f"({sum(1 for f in to_apply if f['status']=='confirmed')} confirmed, "
        f"{sum(1 for f in to_apply if f['status']=='corrected')} corrected)"
    )

    # Korrekturen zusammenstellen
    replacements: List[Tuple[str, str, str]] = []
    for f in to_apply:
        finding_id   = f.get("finding_id", "?")
        original     = str(f.get("original_span") or "").strip()
        status       = f.get("status", "")

        if status == "corrected":
            corrected = str(f.get("human_correction") or "").strip()
            if not corrected:
                print(
                    f"  [SKIP] {finding_id}: status='corrected' aber human_correction leer – "
                    f"übersprungen"
                )
                continue
        else:  # "confirmed"
            corrected = str(f.get("corrected_span") or "").strip()
            if not corrected:
                print(
                    f"  [SKIP] {finding_id}: status='confirmed' aber corrected_span leer – "
                    f"übersprungen (human_correction ausfüllen oder status auf 'corrected' setzen)"
                )
                continue

        replacements.append((finding_id, original, corrected))

    if not replacements:
        print("[WARN] Keine anwendbaren Korrekturen nach Filterung – nichts zu tun.")
        return

    # Ausgabepfad: case_xy_corrected.docx im selben Verzeichnis
    stem     = doc_path.stem
    out_name = f"{stem}_corrected{doc_path.suffix}"
    out_path = doc_path.parent / out_name

    if out_path.exists():
        print(f"[WARN] Ausgabedatei existiert bereits und wird überschrieben: {out_name}")

    print(f"[INFO] Wende {len(replacements)} Korrektur(en) an ...")
    n_ok, n_skip = _apply_replacements_to_docx(doc_path, out_path, replacements)

    sep = "=" * 60
    print(f"\n{sep}")
    print(f"Angewandt: {n_ok} OK | {n_skip} nicht gefunden / übersprungen")
    print(f"Gespeichert: {out_path}")
    if n_ok > 0:
        print(f"\nNÄCHSTER SCHRITT:")
        print(
            f"  python llm_error_detector.py --mode to_gt \\\n"
            f"      --document {doc_path} \\\n"
            f"      --findings {args.findings}"
        )


# ═══════════════════════════════════════════════════════════════════════════════
# MODUS 3: TO_GT
# ═══════════════════════════════════════════════════════════════════════════════

def run_to_gt(args: argparse.Namespace) -> None:
    """
    Konvertiert die vom Enduser bestätigten Findings in eine Ground-Truth-JSONL-Datei,
    die mit evaluate_predictions.py ausgewertet werden kann.

    Das GT-Schema ist identisch mit ground_truth_<case_id>_synthetic.jsonl:
    {
      "case_id":       "case_xy",
      "segment_id":    "case_xy_seg_0003",
      "segment_index": 3,
      "gold_findings": [
        {
          "finding_id":    "GT-case_xy-0001",
          "subclass_id":   "STRUKT_BEFUND_BESCHREIBUNG",
          "change_type_id":"CHANGE_FACHLICH",
          "severity_id":   "HIGH",
          "span_text":     "15,3 km/h",   # original_span (= Fehler im Dokument)
          "correction":    "13,5 km/h",   # was korrekt wäre
          "rationale":     "..."
        }
      ]
    }

    Nur Findings mit status="confirmed" oder status="corrected" werden
    in die GT aufgenommen.

    Der segment_index wird via Volltext-Lookup im Originaldokument verifiziert,
    damit er mit dem Segmentierungsschema der Agenten übereinstimmt
    (analog zu run_validate() im synthetic_error_injector.py).
    """
    print("\n[INFO] Modus: TO_GT")

    doc_path      = _load_doc_arg(args)
    findings_path = _load_findings_arg(args)
    findings      = _read_findings(findings_path)

    accepted = [
        f for f in findings
        if f.get("status") in ("confirmed", "corrected")
    ]

    if not accepted:
        print("[WARN] Keine bestätigten Findings (confirmed/corrected) – keine GT erzeugt.")
        return

    print(
        f"[INFO] Findings gesamt: {len(findings)} | "
        f"für GT: {len(accepted)} "
        f"({sum(1 for f in accepted if f['status']=='confirmed')} confirmed, "
        f"{sum(1 for f in accepted if f['status']=='corrected')} corrected)"
    )

    # Dokument lesen und segmentieren (für segment_index-Verifikation)
    doc_text = normalize_text(read_docx(doc_path))
    segments = split_document_into_segments(
        doc_text,
        target_chars=env_int("SEG_TARGET_CHARS", 1200),
        min_chars=env_int("SEG_MIN_CHARS", 250),
        max_chars=env_int("SEG_MAX_CHARS", 2200),
    )
    print(f"[INFO] Dokumentsegmente: {len(segments)}")

    case_id = (accepted[0].get("case_id") or args.case_id or "case_unknown").strip()

    gt_entries: List[Dict[str, Any]] = []
    finding_counter = 1
    all_ok          = True

    for f in accepted:
        finding_id   = f.get("finding_id", "?")
        original_span = str(f.get("original_span") or "").strip()

        # Korrektur bestimmen
        if f.get("status") == "corrected":
            correction = str(f.get("human_correction") or "").strip()
            if not correction:
                print(
                    f"  [SKIP] {finding_id}: status='corrected' aber human_correction leer"
                )
                all_ok = False
                continue
        else:  # "confirmed"
            correction = str(f.get("corrected_span") or "").strip()
            # correction kann leer sein — GT ist trotzdem gültig (Fehler ohne Korrekturvorschlag)

        if not original_span:
            print(f"  [SKIP] {finding_id}: leerer original_span")
            all_ok = False
            continue

        # segment_index via Dokumenttext verifizieren
        # Zunächst den aus dem Finding übernommenen Index prüfen,
        # dann bei Bedarf neu suchen.
        stored_seg_idx = f.get("segment_index")
        actual_seg_idx: Optional[int] = None

        # Prüfen ob original_span in dem angegebenen Segment liegt
        if stored_seg_idx is not None and 0 <= stored_seg_idx < len(segments):
            if original_span in segments[stored_seg_idx]:
                actual_seg_idx = stored_seg_idx

        # Falls nicht: im gesamten Dokument suchen
        if actual_seg_idx is None:
            actual_seg_idx = _find_segment_for_stelle(original_span, segments)

        if actual_seg_idx is None:
            # Letzter Fallback: Volltext-Check
            if original_span in doc_text:
                print(
                    f"  [WARN] {finding_id}: original_span im Volltext, "
                    f"aber in keinem Segment – Segmentierungsparameter prüfen"
                )
                all_ok = False
                continue
            else:
                print(
                    f"  [FAIL] {finding_id}: original_span "
                    f"'{original_span[:60]}' nicht im Dokument gefunden"
                )
                all_ok = False
                continue

        if stored_seg_idx is not None and actual_seg_idx != stored_seg_idx:
            print(
                f"  [SHIFT] {finding_id}: segment_index verschoben "
                f"{stored_seg_idx} → {actual_seg_idx} "
                f"(GT verwendet korrekten Index {actual_seg_idx})"
            )
        else:
            print(
                f"  [OK]  {finding_id} (seg={actual_seg_idx}): "
                f"'{original_span[:60]}'"
            )

        gt_finding: Dict[str, Any] = {
            "finding_id":    f"GT-{case_id}-{finding_counter:04d}",
            "subclass_id":   f.get("subclass_id", ""),
            "change_type_id": f.get("change_type_id", ""),
            "severity_id":   f.get("severity_id", "MEDIUM"),
            "span_text":     original_span,   # Fehler im Dokument
            "correction":    correction,       # korrekte Version
            "rationale":     f.get("rationale", ""),
            "source":        "llm_detection",
            "detection_finding_id": f.get("finding_id", ""),
        }
        finding_counter += 1

        existing = next(
            (e for e in gt_entries if e["segment_index"] == actual_seg_idx), None
        )
        if existing:
            existing["gold_findings"].append(gt_finding)
        else:
            gt_entries.append({
                "case_id":       case_id,
                "segment_id":    f"{case_id}_seg_{actual_seg_idx:04d}",
                "segment_index": actual_seg_idx,
                "gold_findings": [gt_finding],
            })

    # ── Ground-Truth-JSONL speichern ──────────────────────────────────────────
    gt_dir  = (Path(__file__).parent / "ground_truth").resolve()
    gt_dir.mkdir(parents=True, exist_ok=True)
    gt_path = gt_dir / f"ground_truth_{case_id}_real.jsonl"

    with gt_path.open("w", encoding="utf-8") as f:
        for entry in sorted(gt_entries, key=lambda e: e["segment_index"]):
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")

    total = sum(len(e["gold_findings"]) for e in gt_entries)
    sep   = "=" * 60
    print(f"\n{sep}")
    print(f"Ground Truth: {total} Findings in {len(gt_entries)} Segmenten")
    print(f"Gespeichert:  {gt_path}")
    if not all_ok:
        print("[WARN] Einige Findings konnten nicht zugeordnet werden – Log prüfen.")
    print(f"\nNÄCHSTER SCHRITT (optional – Evaluation gegen Predictions):")
    print(
        f"  python evaluate_predictions.py \\\n"
        f"      --ground_truth_jsonl {gt_path} \\\n"
        f"      --predictions_jsonl  <pfad_zu_predictions.jsonl>"
    )


# ═══════════════════════════════════════════════════════════════════════════════
# HILFSFUNKTIONEN
# ═══════════════════════════════════════════════════════════════════════════════

def _resolve_output_dir(output_dir_str: str) -> Path:
    p = Path(output_dir_str).expanduser()
    d = p if p.is_absolute() else (Path(__file__).parent / p).resolve()
    d.mkdir(parents=True, exist_ok=True)
    return d


def _load_doc_arg(args: argparse.Namespace) -> Path:
    if not args.document:
        sys.exit("[FATAL] --document <pfad> fehlt")
    p = Path(args.document).expanduser().resolve()
    if not p.exists():
        sys.exit(f"[FATAL] Dokument nicht gefunden: {p}")
    if p.suffix.lower() != ".docx":
        sys.exit(f"[FATAL] Nur .docx unterstützt: {p.suffix}")
    print(f"[INFO] Dokument: {p.name}")
    return p


def _load_findings_arg(args: argparse.Namespace) -> Path:
    if not getattr(args, "findings", ""):
        sys.exit("[FATAL] --findings <pfad> fehlt")
    p = Path(args.findings).expanduser().resolve()
    if not p.exists():
        sys.exit(f"[FATAL] Findings-File nicht gefunden: {p}")
    return p


def _read_findings(path: Path) -> List[Dict[str, Any]]:
    with path.open(encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, list):
        sys.exit(f"[FATAL] Findings-File ist keine JSON-Liste: {path}")
    print(f"[INFO] Findings geladen: {len(data)} Einträge aus {path.name}")
    return data


def _print_agent_summary(findings: List[Dict[str, Any]]) -> None:
    """Gibt eine kurze Übersicht der Findings nach Agent aus."""
    counts: Dict[str, int] = {}
    for f in findings:
        agent = f.get("agent", "unknown")
        counts[agent] = counts.get(agent, 0) + 1
    print("Findings nach Agent:")
    for agent, n in sorted(counts.items(), key=lambda x: -x[1]):
        print(f"  {agent:<30} {n:3d}")


# ═══════════════════════════════════════════════════════════════════════════════
# ARGPARSE & MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser(
        description=(
            "LLM-basierte Fehlererkennung in forensischen Gutachten.\n\n"
            "Modi:\n"
            "  detect  – RAG-Agenten laufen lassen, Findings als JSON speichern\n"
            "  apply   – Bestätigte Korrekturen ins Dokument schreiben\n"
            "  to_gt   – Bestätigte Findings als Ground-Truth-JSONL exportieren\n"
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    ap.add_argument(
        "--mode",
        choices=["detect", "apply", "to_gt"],
        required=True,
        help="detect | apply | to_gt",
    )
    ap.add_argument(
        "--document", type=str,
        default=env_str("INJECTION_DOCUMENT_PATH", ""),
        help="Pfad zum Quell-.docx Dokument",
    )
    ap.add_argument(
        "--case_id", type=str,
        default="",
        help="Case-ID. Wenn leer, wird sie aus dem Dokumentnamen abgeleitet.",
    )
    ap.add_argument(
        "--findings", type=str, default="",
        help="[apply / to_gt] Pfad zu findings_<case_id>_<ts>.json",
    )
    ap.add_argument(
        "--output_dir", type=str,
        default=env_str("DETECTION_OUTPUT_DIR", "detection"),
        help="Ausgabeverzeichnis für --mode detect (default: detection)",
    )

    # Taxonomie
    ap.add_argument(
        "--taxonomy_json", type=str,
        default=env_str("TAXONOMY_JSON", "taxonomy.json"),
        help="Pfad zur Taxonomie-JSON",
    )

    # RAG-Store 1 (Normbasis)
    # Primär RAG_* für den GUI-Wrapper; Fallback auf die Namen aus rag_answer_reference_facts.py.
    ap.add_argument("--embeddings",  type=str, default=env_str("RAG_EMBEDDINGS", env_str("EMBED_OUT_NPZ", "embeddings.npz")))
    ap.add_argument("--index",       type=str, default=env_str("RAG_INDEX", env_str("EMBED_OUT_INDEX", "index.jsonl")))
    ap.add_argument("--prepared",    type=str, default=env_str("RAG_PREPARED", env_str("OUT_JSONL", "prepared.jsonl")))

    # RAG-Store 2 (Zusatzmaterialien, optional)
    ap.add_argument("--embeddings2", type=str, default=env_str("RAG_EMBEDDINGS2", env_str("EMBED_OUT_NPZ2", "")))
    ap.add_argument("--index2",      type=str, default=env_str("RAG_INDEX2", env_str("EMBED_OUT_INDEX2", "")))
    ap.add_argument("--prepared2",   type=str, default=env_str("RAG_PREPARED2", env_str("OUT_JSONL2", "")))

    # Embedding-Konfiguration
    ap.add_argument(
        "--embed_model", type=str,
        default=env_str("EMBED_MODEL", "intfloat/multilingual-e5-large"),
    )
    ap.add_argument(
        "--embed_device", type=str,
        default=env_str("EMBED_DEVICE", "auto"),
    )
    ap.add_argument(
        "--query_max_length", type=int,
        default=env_int("QUERY_MAX_LENGTH", 256),
    )

    # Retrieval-Konfiguration (werden an build_segment_evidences weitergegeben)
    ap.add_argument("--top_k",              type=int,   default=env_int("TOP_K", 10))
    ap.add_argument("--rules_top_k",        type=int,   default=env_int("RULES_TOP_K", 8))
    ap.add_argument("--material_top_k",     type=int,   default=env_int("MATERIAL_TOP_K", 8))
    ap.add_argument("--context_max_chars",  type=int,   default=env_int("CONTEXT_MAX_CHARS", 12000))
    ap.add_argument("--multi_query_count",  type=int,   default=env_int("MULTI_QUERY_COUNT", 4))
    ap.add_argument("--mmr_lambda",         type=float, default=float(env_str("MMR_LAMBDA", "0.75")))
    ap.add_argument("--max_per_source",     type=int,   default=env_int("MAX_PER_SOURCE", 2))
    ap.add_argument("--per_segment_candidate_k",   type=int, default=env_int("PER_SEGMENT_CANDIDATE_K", 24))
    ap.add_argument("--per_segment_rules_top_k",   type=int, default=env_int("PER_SEGMENT_RULES_TOP_K", 8))
    ap.add_argument("--per_segment_material_top_k",type=int, default=env_int("PER_SEGMENT_MATERIAL_TOP_K", 8))

    # Vision (optional)
    ap.add_argument("--vision_model",   type=str, default=env_str("VISION_MODEL", ""))
    ap.add_argument("--vision_workers", type=int, default=env_int("VISION_WORKERS", 3))
    ap.add_argument("--vision_timeout_s", type=int, default=env_int("VISION_TIMEOUT_S", 180))

    # Referenzfakten-Agent
    ap.add_argument(
        "--reference_facts_schema", type=str,
        default=env_str("REFERENCE_FACTS_SCHEMA", "schema/reference_facts_schema.json"),
    )
    ap.add_argument(
        "--reference_facts_context_chars", type=int,
        default=env_int("REFERENCE_FACTS_CONTEXT_CHARS", 4000),
    )

    # Output-Artefakte der Hauptpipeline / Debugging
    ap.add_argument(
        "--save_predictions_jsonl",
        type=str,
        default="",
        help="Optionaler Pfad für predictions_<case_id>.jsonl. Wenn leer: predictions/predictions_<case_id>.jsonl",
    )
    ap.add_argument(
        "--save_reference_facts_json",
        type=str,
        default="",
        help="Optionaler Pfad für reference_facts_<case_id>.json. Wenn leer: reference_facts/reference_facts_<case_id>.json",
    )
    ap.add_argument(
        "--save_report_json",
        type=str,
        default="",
        help="Optionaler Pfad für vollständiges report_<case_id>.json. Wenn leer: reports/report_<case_id>.json",
    )

    # Print-Flags (für Debugging; im GUI-Detect per DETECTION_PRINT_* standardmässig aktiv)
    ap.add_argument("--print_sources",  action="store_true")
    ap.add_argument("--print_context",  action="store_true")
    ap.add_argument("--print_reference_facts", action="store_true")

    args = ap.parse_args()

    # case_id: explizit → Dokument-Stem → Findings-File → ENV → Fallback
    # Wichtig für die GUI: Wenn case_14.docx ausgewählt wird, darf ein altes
    # CASE_ID=case_06 aus .env nicht den neuen Fall überschreiben.
    if str(args.case_id or "").strip():
        args.case_id = str(args.case_id).strip()
    elif args.document:
        args.case_id = Path(args.document).stem
    elif args.findings:
        try:
            data = json.loads(Path(args.findings).read_text(encoding="utf-8"))
            args.case_id = (data[0].get("case_id") or "case_unknown") if data else "case_unknown"
        except Exception:
            args.case_id = "case_unknown"
    else:
        args.case_id = env_str("CASE_ID", "") or "case_unknown"

    return args


def main() -> None:
    t0 = time.perf_counter()
    print(
        f"[INFO] llm_error_detector.py – "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    args = parse_args()

    if args.mode == "detect":
        if not args.document:
            sys.exit("[FATAL] --mode detect erfordert --document")
        if not args.embeddings or not args.index or not args.prepared:
            sys.exit(
                "[FATAL] --mode detect erfordert RAG-Store-Dateien "
                "(--embeddings/--index/--prepared oder RAG_* bzw. EMBED_* in .env)"
            )
        run_detect(args)

    elif args.mode == "apply":
        if not args.document:
            sys.exit("[FATAL] --mode apply erfordert --document")
        if not args.findings:
            sys.exit("[FATAL] --mode apply erfordert --findings")
        run_apply(args)

    elif args.mode == "to_gt":
        if not args.document:
            sys.exit("[FATAL] --mode to_gt erfordert --document")
        if not args.findings:
            sys.exit("[FATAL] --mode to_gt erfordert --findings")
        run_to_gt(args)

    elapsed = time.perf_counter() - t0
    print(f"\n[INFO] Fertig in {elapsed:.1f}s ({elapsed/60:.1f} min)")


if __name__ == "__main__":
    main()
